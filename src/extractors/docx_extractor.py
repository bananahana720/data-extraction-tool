"""
DOCX Extractor - Microsoft Word Document Extraction

This is a SPIKE implementation focused on getting working text extraction
from Word documents. It extracts paragraphs with basic metadata.

Current Scope:
- Extract text paragraphs
- Basic position tracking (sequence)
- Document metadata
- Error handling
- Tables (DOCX-TABLE-001) ✓

Not Yet Implemented:
- Images (DOCX-IMAGE-001)
- Headers/footers (DOCX-HEADER-001)
- Styles/formatting details (DOCX-STYLE-001)
- Lists (DOCX-LIST-001)
- Footnotes/comments (DOCX-META-001)
"""

from pathlib import Path
from typing import Optional, Union
from datetime import datetime, timezone
import hashlib
import logging

try:
    from docx import Document
    from docx.oxml.exceptions import InvalidXmlError
except ImportError:
    raise ImportError(
        "python-docx is required for DocxExtractor. " "Install with: pip install python-docx"
    )

import sys

sys.path.insert(0, str(Path(__file__).parent.parent))

from core import (
    BaseExtractor,
    ContentBlock,
    ContentType,
    DocumentMetadata,
    ExtractionResult,
    Position,
    TableMetadata,
)

# Import infrastructure components
try:
    from infrastructure import (
        ConfigManager,
        get_logger,
        timed,
        timer,
        ErrorHandler,
        ProgressTracker,
        RecoveryAction,
    )

    INFRASTRUCTURE_AVAILABLE = True
except ImportError:
    INFRASTRUCTURE_AVAILABLE = False

    # Fallback: Create dummy decorator for compatibility
    def timed(logger, level=logging.INFO):
        def decorator(func):
            return func

        return decorator


class DocxExtractor(BaseExtractor):
    """
    Extracts text content from Microsoft Word (.docx) files.

    This is a SPIKE implementation focused on paragraph extraction.
    Uses python-docx library for document parsing.

    Design Notes:
    - Only extracts text paragraphs in this version
    - Detects headings based on style names
    - Handles basic errors gracefully
    - Returns partial results on non-fatal errors

    Example:
        >>> extractor = DocxExtractor()
        >>> result = extractor.extract(Path("document.docx"))
        >>> if result.success:
        ...     for block in result.content_blocks:
        ...         print(block.content)
    """

    def __init__(self, config: Optional[Union[dict, object]] = None):
        """
        Initialize DOCX extractor with optional configuration.

        Args:
            config: Configuration options (dict or ConfigManager):
                - max_paragraph_length: Max characters per paragraph (default: None)
                - skip_empty: Skip empty paragraphs (default: True)
                - extract_styles: Include style information (default: True)
        """
        super().__init__(config if isinstance(config, dict) or config is None else {})

        # Store original config for infrastructure access
        # Note: Use class name check to avoid import path mismatches
        is_config_manager = (
            INFRASTRUCTURE_AVAILABLE
            and hasattr(config, "__class__")
            and config.__class__.__name__ == "ConfigManager"
        )
        self._config_manager = config if is_config_manager else None

        # Initialize infrastructure components
        if INFRASTRUCTURE_AVAILABLE:
            self.logger = get_logger(__name__)
            self.error_handler = ErrorHandler()
        else:
            self.logger = logging.getLogger(__name__)
            self.error_handler = None

        # Load configuration
        if self._config_manager:
            # Use ConfigManager - get from extractor section
            extractor_config = self._config_manager.get_section("extractors.docx", default={})
            # Note: Use explicit check for None to handle False values correctly
            self.max_paragraph_length = extractor_config.get("max_paragraph_length")
            skip_empty_val = extractor_config.get("skip_empty")
            self.skip_empty = skip_empty_val if skip_empty_val is not None else True
            extract_styles_val = extractor_config.get("extract_styles")
            self.extract_styles = extract_styles_val if extract_styles_val is not None else True
        elif isinstance(config, dict):
            # Use dict config (backward compatibility)
            self.max_paragraph_length = config.get("max_paragraph_length", None)
            self.skip_empty = config.get("skip_empty", True)
            self.extract_styles = config.get("extract_styles", True)
        else:
            # Use defaults
            self.max_paragraph_length = None
            self.skip_empty = True
            self.extract_styles = True

    def supports_format(self, file_path: Path) -> bool:
        """
        Check if file is a DOCX file.

        Args:
            file_path: Path to check

        Returns:
            True if file has .docx extension
        """
        return file_path.suffix.lower() == ".docx"

    def get_supported_extensions(self) -> list[str]:
        """Return supported file extensions."""
        return [".docx"]

    def get_format_name(self) -> str:
        """Return human-readable format name."""
        return "Microsoft Word"

    def extract(self, file_path: Path) -> ExtractionResult:
        """
        Extract content from DOCX file.

        Strategy:
        1. Validate file exists and is accessible
        2. Open document with python-docx
        3. Extract paragraphs sequentially
        4. Detect content types (heading vs paragraph)
        5. Generate document metadata
        6. Return structured result

        Args:
            file_path: Path to DOCX file

        Returns:
            ExtractionResult with content blocks and metadata

        Note:
            - Returns success=False for file-level errors
            - Returns partial results if some paragraphs fail
            - Logs warnings for recoverable issues
        """
        import time

        start_time = time.time()

        # Log extraction start
        if INFRASTRUCTURE_AVAILABLE:
            self.logger.info("Starting DOCX extraction", extra={"file": str(file_path)})

        errors = []
        warnings = []
        content_blocks = []

        # Step 1: Validate file
        is_valid, validation_errors = self.validate_file(file_path)
        if not is_valid:
            # Use error handler if available
            if self.error_handler:
                error = self.error_handler.create_error("E001", file_path=str(file_path))
                errors.append(self.error_handler.format_for_user(error))
                if INFRASTRUCTURE_AVAILABLE:
                    self.logger.error(
                        "File validation failed",
                        extra={"file": str(file_path), "errors": validation_errors},
                    )
            else:
                errors.extend(validation_errors)

            return ExtractionResult(
                success=False,
                errors=tuple(errors),
                document_metadata=DocumentMetadata(
                    source_file=file_path,
                    file_format="docx",
                ),
            )

        try:
            # Step 2: Open document
            try:
                doc = Document(file_path)
            except Exception as e:
                return ExtractionResult(
                    success=False,
                    errors=(f"Failed to open DOCX file: {str(e)}",),
                    document_metadata=DocumentMetadata(
                        source_file=file_path,
                        file_format="docx",
                    ),
                )

            # Step 3: Extract paragraphs
            paragraph_count = 0
            for idx, paragraph in enumerate(doc.paragraphs):
                # Get text content
                text = paragraph.text.strip()

                # Skip empty paragraphs if configured
                if not text and self.skip_empty:
                    continue

                # Check length limit
                if self.max_paragraph_length and len(text) > self.max_paragraph_length:
                    warnings.append(
                        f"Paragraph {idx} truncated from {len(text)} to "
                        f"{self.max_paragraph_length} characters"
                    )
                    text = text[: self.max_paragraph_length] + "..."

                # Step 4: Detect content type
                block_type = self._detect_content_type(paragraph)

                # Build metadata
                metadata = {
                    "paragraph_index": idx,
                    "char_count": len(text),
                    "word_count": len(text.split()) if text else 0,
                }

                # Add style information if configured
                if self.extract_styles and paragraph.style:
                    metadata["style_name"] = paragraph.style.name

                # Create content block
                block = ContentBlock(
                    block_type=block_type,
                    content=text,
                    raw_content=paragraph.text,  # Unstripped original
                    position=Position(sequence_index=paragraph_count),
                    confidence=1.0,  # High confidence for native format
                    metadata=metadata,
                )

                content_blocks.append(block)
                paragraph_count += 1

            # Check if we got any content
            if not content_blocks:
                warnings.append("No content extracted from document")

            # Step 4.5: Extract tables
            tables = []
            for table_idx, table in enumerate(doc.tables):
                try:
                    table_metadata = self._extract_table(table, table_idx)
                    tables.append(table_metadata)
                except Exception as e:
                    warnings.append(f"Failed to extract table {table_idx}: {str(e)}")

            # Step 5: Generate document metadata
            doc_metadata = self._extract_document_metadata(file_path, doc)

            # Update statistics
            total_chars = sum(len(b.content) for b in content_blocks)
            total_words = sum(len(b.content.split()) for b in content_blocks)

            doc_metadata = DocumentMetadata(
                source_file=doc_metadata.source_file,
                file_format=doc_metadata.file_format,
                file_size_bytes=doc_metadata.file_size_bytes,
                file_hash=doc_metadata.file_hash,
                title=doc_metadata.title,
                author=doc_metadata.author,
                created_date=doc_metadata.created_date,
                modified_date=doc_metadata.modified_date,
                subject=doc_metadata.subject,
                keywords=doc_metadata.keywords,
                word_count=total_words,
                character_count=total_chars,
                table_count=len(tables),
                extracted_at=doc_metadata.extracted_at,
                extractor_version="0.1.0-spike",
            )

            # Step 6: Log completion and return result
            duration = time.time() - start_time
            if INFRASTRUCTURE_AVAILABLE:
                self.logger.info(
                    "DOCX extraction complete",
                    extra={
                        "file": str(file_path),
                        "blocks": len(content_blocks),
                        "words": total_words,
                        "duration_seconds": round(duration, 3),
                    },
                )

            return ExtractionResult(
                content_blocks=tuple(content_blocks),
                document_metadata=doc_metadata,
                tables=tuple(tables),
                success=True,
                warnings=tuple(warnings),
            )

        except InvalidXmlError as e:
            if self.error_handler:
                error = self.error_handler.create_error(
                    "E110", file_path=str(file_path), original_exception=e
                )
                errors.append(self.error_handler.format_for_user(error))
            else:
                errors.append(f"DOCX structure error: {str(e)}")

            if INFRASTRUCTURE_AVAILABLE:
                self.logger.error(
                    "DOCX structure error", extra={"file": str(file_path), "error": str(e)}
                )

        except PermissionError as e:
            if self.error_handler:
                error = self.error_handler.create_error(
                    "E500", file_path=str(file_path), original_exception=e
                )
                errors.append(self.error_handler.format_for_user(error))
            else:
                errors.append(f"Permission denied reading file: {file_path}")

            if INFRASTRUCTURE_AVAILABLE:
                self.logger.error("Permission denied", extra={"file": str(file_path)})

        except Exception as e:
            if self.error_handler:
                error = self.error_handler.create_error(
                    "E100", file_path=str(file_path), original_exception=e
                )
                errors.append(self.error_handler.format_for_user(error))
            else:
                errors.append(f"Unexpected error during extraction: {str(e)}")

            if INFRASTRUCTURE_AVAILABLE:
                self.logger.error(
                    "Extraction failed", extra={"file": str(file_path), "error": str(e)}
                )

        # Return failed result
        return ExtractionResult(
            success=False,
            errors=tuple(errors),
            warnings=tuple(warnings),
            document_metadata=DocumentMetadata(
                source_file=file_path,
                file_format="docx",
            ),
        )

    def _detect_content_type(self, paragraph) -> ContentType:
        """
        Detect content type based on paragraph style.

        Uses Word's built-in style names to classify content.

        Args:
            paragraph: python-docx Paragraph object

        Returns:
            ContentType enum value
        """
        if not paragraph.style:
            return ContentType.PARAGRAPH

        style_name = paragraph.style.name.lower()

        # Check for heading styles
        if "heading" in style_name:
            return ContentType.HEADING

        # Check for list styles
        if "list" in style_name:
            return ContentType.LIST_ITEM

        # Check for quote/block quote
        if "quote" in style_name or "block" in style_name:
            return ContentType.QUOTE

        # Check for code
        if "code" in style_name or "source" in style_name:
            return ContentType.CODE

        # Default to paragraph
        return ContentType.PARAGRAPH

    def _extract_document_metadata(self, file_path: Path, doc: Document) -> DocumentMetadata:
        """
        Extract document-level metadata from DOCX file.

        Extracts both file system metadata and embedded document properties.

        Args:
            file_path: Path to file
            doc: python-docx Document object

        Returns:
            DocumentMetadata with available properties
        """
        # File system metadata
        file_stat = file_path.stat()
        file_size = file_stat.st_size

        # Generate file hash for deduplication
        file_hash = self._compute_file_hash(file_path)

        # Extract core properties (document metadata)
        core_props = doc.core_properties

        # Parse dates (may be None)
        created = core_props.created
        modified = core_props.modified

        # Parse keywords (may be None or string)
        keywords = ()
        if core_props.keywords:
            keywords = tuple(k.strip() for k in core_props.keywords.split(",") if k.strip())

        return DocumentMetadata(
            source_file=file_path,
            file_format="docx",
            file_size_bytes=file_size,
            file_hash=file_hash,
            title=core_props.title or None,
            author=core_props.author or None,
            created_date=created if isinstance(created, datetime) else None,
            modified_date=modified if isinstance(modified, datetime) else None,
            subject=core_props.subject or None,
            keywords=keywords,
            extracted_at=datetime.now(timezone.utc),
        )

    def _compute_file_hash(self, file_path: Path) -> str:
        """
        Compute SHA256 hash of file for deduplication.

        Args:
            file_path: Path to file

        Returns:
            Hex string of SHA256 hash
        """
        sha256 = hashlib.sha256()

        # Read in chunks to handle large files
        with open(file_path, "rb") as f:
            for chunk in iter(lambda: f.read(8192), b""):
                sha256.update(chunk)

        return sha256.hexdigest()

    def _extract_table(self, table, table_idx: int) -> TableMetadata:
        """
        Extract table data from a python-docx Table object.

        Args:
            table: python-docx Table object
            table_idx: Index of table in document

        Returns:
            TableMetadata with table structure and content
        """
        # Get table dimensions
        num_rows = len(table.rows)
        num_columns = len(table.columns) if table.rows else 0

        # Extract all cells
        cells = []
        for row in table.rows:
            row_cells = []
            for cell in row.cells:
                # Get cell text, strip whitespace
                cell_text = cell.text.strip()
                row_cells.append(cell_text)
            cells.append(tuple(row_cells))

        # Detect if first row is header (heuristic: check if styled differently)
        has_header = False
        header_row = None
        if cells:
            # Simple heuristic: assume first row is header if it exists
            # More sophisticated detection could check cell styles
            has_header = True
            header_row = cells[0]

        return TableMetadata(
            num_rows=num_rows,
            num_columns=num_columns,
            has_header=has_header,
            header_row=header_row,
            cells=tuple(cells),
        )
