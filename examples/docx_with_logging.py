"""
Example showing DocxExtractor with logging framework integration.

This demonstrates how Wave 2 Agent 4 (DocxExtractorRefactor) can add
logging to the existing DocxExtractor.
"""

import sys
from pathlib import Path

# Add src to path
sys.path.insert(0, str(Path(__file__).parent.parent / "src"))

from infrastructure import get_logger, timed, timer
from extractors.docx_extractor import DocxExtractor


# Create logger at module level
logger = get_logger(__name__, console=True, json_format=False)


def main():
    """Demonstrate DocxExtractor with logging."""

    print("=" * 70)
    print("DocxExtractor with Logging Framework")
    print("=" * 70)

    # Get sample file
    sample_file = Path(__file__).parent.parent / "tests" / "fixtures" / "sample.docx"

    if not sample_file.exists():
        print(f"\nError: Sample file not found: {sample_file}")
        print("Creating a test document instead...")

        # Try to create a simple test document
        try:
            from docx import Document
            test_doc_path = Path("test_document.docx")
            doc = Document()
            doc.add_heading("Test Document", 0)
            doc.add_paragraph("This is a test paragraph.")
            doc.add_paragraph("This is another paragraph with more content.")
            doc.save(test_doc_path)
            sample_file = test_doc_path
            print(f"Created test document: {test_doc_path}")
        except ImportError:
            print("\nError: python-docx not installed. Cannot create test document.")
            print("Install with: pip install python-docx")
            return

    logger.info("Starting extraction demo", extra={"file": str(sample_file)})

    # Create extractor
    extractor = DocxExtractor(config={
        "skip_empty": True,
        "extract_styles": True
    })

    # Extract with timing
    with timer(logger, "docx_extraction"):
        result = extractor.extract(sample_file)

    # Log results
    if result.success:
        logger.info(
            "Extraction successful",
            extra={
                "file": str(sample_file),
                "blocks": len(result.content_blocks),
                "has_warnings": len(result.warnings) > 0,
                "word_count": result.document_metadata.word_count,
                "character_count": result.document_metadata.character_count
            }
        )

        print(f"\nExtraction Results:")
        print(f"  - Blocks extracted: {len(result.content_blocks)}")
        print(f"  - Word count: {result.document_metadata.word_count}")
        print(f"  - Character count: {result.document_metadata.character_count}")

        if result.warnings:
            logger.warning("Extraction had warnings", extra={"count": len(result.warnings)})
            print(f"\nWarnings ({len(result.warnings)}):")
            for warning in result.warnings:
                print(f"  - {warning}")

        print("\nFirst 3 content blocks:")
        for idx, block in enumerate(result.content_blocks[:3]):
            logger.debug(
                "Content block",
                extra={
                    "index": idx,
                    "type": block.block_type.value,
                    "length": len(block.content)
                }
            )
            print(f"\n  Block {idx + 1}:")
            print(f"    Type: {block.block_type.value}")
            print(f"    Content: {block.content[:100]}...")
    else:
        logger.error(
            "Extraction failed",
            extra={
                "file": str(sample_file),
                "errors": result.errors
            }
        )
        print(f"\nExtraction failed!")
        print(f"Errors:")
        for error in result.errors:
            print(f"  - {error}")

    print("\n" + "=" * 70)
    print("Demo complete!")
    print("=" * 70)
    print("\nThis example shows how logging can be added to extractors.")
    print("Wave 2 Agent 4 will refactor DocxExtractor to include logging.")


if __name__ == "__main__":
    main()
