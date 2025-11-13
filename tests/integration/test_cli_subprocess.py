"""
Subprocess-based CLI Integration Tests.

Tests the installed data-extract CLI command via subprocess.run() to validate:
- Entry point registration and PATH availability
- Real stdout/stderr output (not simulated)
- Exit codes and error handling
- Shell integration and environment variable handling
- Installed package behavior (vs CliRunner simulation)

These tests complement CliRunner tests by validating the actual installed
console script as users will experience it.

Test Categories:
- Basic commands (help, version)
- Extract command with real files
- Batch processing
- Error handling and exit codes
- Configuration file integration
- Environment variable handling
"""

import json
import subprocess

import pytest

# Mark all tests in this module with 'cli' and 'subprocess' markers
pytestmark = [pytest.mark.cli, pytest.mark.subprocess]


class TestCLIInstallation:
    """Test that CLI is properly installed and accessible."""

    def test_cli_command_exists(self):
        """Verify data-extract command is on PATH."""
        result = subprocess.run(
            ["data-extract", "--help"],
            capture_output=True,
            text=True,
            timeout=10,
        )
        # Command should execute without error
        assert result.returncode == 0, f"CLI not found on PATH: {result.stderr}"

    def test_cli_shows_help(self):
        """Verify --help flag displays usage information."""
        result = subprocess.run(
            ["data-extract", "--help"],
            capture_output=True,
            text=True,
            timeout=10,
        )
        assert result.returncode == 0
        assert "Data Extraction Tool" in result.stdout
        assert "Usage:" in result.stdout or "Commands:" in result.stdout

    def test_cli_version_command(self):
        """Verify version command displays version info."""
        result = subprocess.run(
            ["data-extract", "version"],
            capture_output=True,
            text=True,
            timeout=10,
        )
        assert result.returncode == 0
        assert "version" in result.stdout.lower()
        # Should contain a semantic version number (x.y.z)
        assert any(char.isdigit() for char in result.stdout)

    def test_cli_version_verbose(self):
        """Verify version --verbose shows detailed info."""
        result = subprocess.run(
            ["data-extract", "version", "--verbose"],
            capture_output=True,
            text=True,
            timeout=10,
        )
        assert result.returncode == 0
        assert "Python version" in result.stdout or "Platform" in result.stdout


class TestExtractCommandSubprocess:
    """Test extract command via subprocess."""

    def test_extract_help(self):
        """Verify extract --help works."""
        result = subprocess.run(
            ["data-extract", "extract", "--help"],
            capture_output=True,
            text=True,
            timeout=10,
        )
        assert result.returncode == 0
        assert "extract" in result.stdout.lower()

    def test_extract_missing_file_error(self, tmp_path):
        """Verify graceful error on missing file."""
        nonexistent = tmp_path / "nonexistent.pdf"
        result = subprocess.run(
            ["data-extract", "extract", str(nonexistent)],
            capture_output=True,
            text=True,
            timeout=10,
        )
        # Should fail with non-zero exit code
        assert result.returncode != 0
        # Error message should mention file not found
        output = result.stdout + result.stderr
        assert (
            "not found" in output.lower()
            or "does not exist" in output.lower()
            or "no such file" in output.lower()
        )

    def test_extract_docx_to_json(self, tmp_path):
        """End-to-end: Extract DOCX to JSON via subprocess."""
        # Create minimal DOCX fixture
        from docx import Document

        input_file = tmp_path / "test.docx"
        doc = Document()
        doc.add_paragraph("Test content for subprocess validation.")
        doc.save(input_file)

        output_file = tmp_path / "output.json"

        result = subprocess.run(
            [
                "data-extract",
                "extract",
                str(input_file),
                "--output",
                str(output_file),
                "--format",
                "json",
            ],
            capture_output=True,
            text=True,
            timeout=30,
        )

        # Verify success
        assert result.returncode == 0, f"Extract failed: {result.stderr}"

        # Verify output file created
        assert output_file.exists(), "Output file not created"

        # Verify valid JSON
        with open(output_file, encoding="utf-8") as f:
            data = json.load(f)
            assert "content" in data or "blocks" in data or "text" in data

    def test_extract_with_quiet_flag(self, tmp_path):
        """Verify --quiet suppresses progress output."""
        from docx import Document

        input_file = tmp_path / "test.docx"
        doc = Document()
        doc.add_paragraph("Test content.")
        doc.save(input_file)

        output_file = tmp_path / "output.json"

        result = subprocess.run(
            [
                "data-extract",
                "--quiet",
                "extract",
                str(input_file),
                "--output",
                str(output_file),
            ],
            capture_output=True,
            text=True,
            timeout=30,
        )

        assert result.returncode == 0
        # Quiet mode should produce minimal stdout
        assert len(result.stdout) < 100 or result.stdout.strip() == ""

    def test_extract_with_verbose_flag(self, tmp_path):
        """Verify --verbose shows detailed output."""
        from docx import Document

        input_file = tmp_path / "test.docx"
        doc = Document()
        doc.add_paragraph("Test content.")
        doc.save(input_file)

        output_file = tmp_path / "output.json"

        result = subprocess.run(
            [
                "data-extract",
                "--verbose",
                "extract",
                str(input_file),
                "--output",
                str(output_file),
            ],
            capture_output=True,
            text=True,
            timeout=30,
        )

        assert result.returncode == 0
        # Verbose mode should show additional details
        # (exact content depends on implementation)


class TestBatchCommandSubprocess:
    """Test batch command via subprocess."""

    def test_batch_help(self):
        """Verify batch --help works."""
        result = subprocess.run(
            ["data-extract", "batch", "--help"],
            capture_output=True,
            text=True,
            timeout=10,
        )
        assert result.returncode == 0
        assert "batch" in result.stdout.lower()

    def test_batch_requires_output_argument(self, tmp_path):
        """Verify batch command requires --output argument."""
        result = subprocess.run(
            ["data-extract", "batch", str(tmp_path)],
            capture_output=True,
            text=True,
            timeout=10,
        )
        # Should fail without --output
        assert result.returncode != 0
        output = result.stdout + result.stderr
        assert "output" in output.lower() or "required" in output.lower()

    def test_batch_process_multiple_files(self, tmp_path):
        """End-to-end: Batch process multiple files via subprocess."""
        from docx import Document

        # Create test files
        input_dir = tmp_path / "inputs"
        input_dir.mkdir()

        for i in range(3):
            file_path = input_dir / f"doc{i}.docx"
            doc = Document()
            doc.add_paragraph(f"Content for document {i}")
            doc.save(file_path)

        output_dir = tmp_path / "outputs"

        result = subprocess.run(
            [
                "data-extract",
                "batch",
                str(input_dir),
                "--output",
                str(output_dir),
                "--format",
                "json",
            ],
            capture_output=True,
            text=True,
            timeout=60,
        )

        # Verify success
        assert result.returncode == 0, f"Batch failed: {result.stderr}"

        # Verify output directory created
        assert output_dir.exists(), "Output directory not created"

        # Verify output files created
        output_files = list(output_dir.glob("*.json"))
        assert len(output_files) >= 3, f"Expected 3+ output files, got {len(output_files)}"

    def test_batch_with_pattern_filter(self, tmp_path):
        """Verify --pattern flag filters files correctly."""
        from docx import Document

        # Create mixed file types
        input_dir = tmp_path / "inputs"
        input_dir.mkdir()

        # Create DOCX files
        for i in range(2):
            file_path = input_dir / f"doc{i}.docx"
            doc = Document()
            doc.add_paragraph(f"DOCX content {i}")
            doc.save(file_path)

        # Create TXT files (should be filtered out with pattern)
        for i in range(2):
            file_path = input_dir / f"text{i}.txt"
            file_path.write_text(f"TXT content {i}")

        output_dir = tmp_path / "outputs"

        result = subprocess.run(
            [
                "data-extract",
                "batch",
                str(input_dir),
                "--pattern",
                "*.docx",
                "--output",
                str(output_dir),
            ],
            capture_output=True,
            text=True,
            timeout=60,
        )

        assert result.returncode == 0

        # Should only process DOCX files
        output_files = list(output_dir.glob("*"))
        # Expecting 2 DOCX outputs (pattern filters out TXT)
        assert len(output_files) == 2, f"Expected 2 outputs, got {len(output_files)}"


class TestConfigCommandSubprocess:
    """Test config command via subprocess."""

    def test_config_help(self):
        """Verify config --help works."""
        result = subprocess.run(
            ["data-extract", "config", "--help"],
            capture_output=True,
            text=True,
            timeout=10,
        )
        assert result.returncode == 0
        assert "config" in result.stdout.lower()

    def test_config_validate_missing_file(self, tmp_path):
        """Verify config validate fails gracefully on missing file."""
        nonexistent_config = tmp_path / "nonexistent-config.yaml"

        result = subprocess.run(
            [
                "data-extract",
                "--config",
                str(nonexistent_config),
                "config",
                "validate",
            ],
            capture_output=True,
            text=True,
            timeout=10,
        )

        # Should fail with non-zero exit code
        assert result.returncode != 0

    def test_config_show_with_valid_file(self, tmp_path):
        """Verify config show displays config contents."""
        import yaml

        config_file = tmp_path / "test-config.yaml"
        config_data = {
            "pipeline": {"max_workers": 4},
            "output": {"default_format": "json"},
        }
        with open(config_file, "w") as f:
            yaml.dump(config_data, f)

        result = subprocess.run(
            ["data-extract", "--config", str(config_file), "config", "show"],
            capture_output=True,
            text=True,
            timeout=10,
        )

        assert result.returncode == 0
        # Config contents should be displayed
        assert "pipeline" in result.stdout or "max_workers" in result.stdout


class TestErrorHandling:
    """Test error handling and exit codes."""

    def test_invalid_command_fails(self):
        """Verify invalid command returns non-zero exit code."""
        result = subprocess.run(
            ["data-extract", "invalid-command-xyz"],
            capture_output=True,
            text=True,
            timeout=10,
        )
        assert result.returncode != 0

    def test_unsupported_format_error(self, tmp_path):
        """Verify unsupported file format fails gracefully."""
        unsupported_file = tmp_path / "unsupported.xyz"
        unsupported_file.write_text("Unsupported format content")

        result = subprocess.run(
            ["data-extract", "extract", str(unsupported_file)],
            capture_output=True,
            text=True,
            timeout=10,
        )

        assert result.returncode != 0
        output = result.stdout + result.stderr
        assert "unsupported" in output.lower() or "format" in output.lower()

    def test_keyboard_interrupt_handling(self, tmp_path):
        """Verify Ctrl+C handling (limited test - can't easily simulate SIGINT)."""
        # This is a placeholder - real SIGINT testing requires tmux-cli or pexpect
        # Subprocess tests can't easily send signals mid-execution
        # For now, just verify the signal handler is registered
        pass

    def test_exit_code_success(self, tmp_path):
        """Verify successful operation returns exit code 0."""
        from docx import Document

        input_file = tmp_path / "test.docx"
        doc = Document()
        doc.add_paragraph("Test")
        doc.save(input_file)

        result = subprocess.run(
            ["data-extract", "extract", str(input_file)],
            capture_output=True,
            text=True,
            timeout=30,
        )

        assert result.returncode == 0

    def test_exit_code_failure(self, tmp_path):
        """Verify failed operation returns non-zero exit code."""
        nonexistent = tmp_path / "nonexistent.pdf"

        result = subprocess.run(
            ["data-extract", "extract", str(nonexistent)],
            capture_output=True,
            text=True,
            timeout=10,
        )

        assert result.returncode != 0


class TestEnvironmentIntegration:
    """Test environment variable and PATH integration."""

    def test_cli_uses_system_python(self):
        """Verify CLI uses correct Python interpreter."""
        result = subprocess.run(
            ["data-extract", "version", "--verbose"],
            capture_output=True,
            text=True,
            timeout=10,
        )

        if result.returncode == 0 and "Python version" in result.stdout:
            # Should show Python 3.12+ (per project requirements)
            assert "3.1" in result.stdout  # Matches 3.12, 3.13, etc.

    def test_cli_encoding_utf8(self, tmp_path):
        """Verify CLI handles UTF-8 output correctly."""
        from docx import Document

        input_file = tmp_path / "test-unicode.docx"
        doc = Document()
        doc.add_paragraph("Unicode test: café, naïve, 日本語, 中文")
        doc.save(input_file)

        output_file = tmp_path / "output.json"

        result = subprocess.run(
            [
                "data-extract",
                "extract",
                str(input_file),
                "--output",
                str(output_file),
            ],
            capture_output=True,
            text=True,
            timeout=30,
        )

        # Should not crash on Unicode
        assert result.returncode == 0

        # Output file should contain Unicode
        content = output_file.read_text(encoding="utf-8")
        assert "café" in content or "Unicode" in content


class TestRealWorldScenarios:
    """Test realistic user workflows."""

    @pytest.mark.slow
    def test_full_workflow_extract_to_json(self, tmp_path):
        """Simulate real user workflow: extract document to JSON."""
        from docx import Document

        # User creates document
        input_file = tmp_path / "audit-report.docx"
        doc = Document()
        doc.add_heading("Audit Report", level=1)
        doc.add_paragraph("Executive Summary: All controls are effective.")
        doc.add_paragraph("Findings: No major issues identified.")
        doc.save(input_file)

        # User extracts to JSON
        output_file = tmp_path / "audit-report.json"

        result = subprocess.run(
            [
                "data-extract",
                "extract",
                str(input_file),
                "--output",
                str(output_file),
                "--format",
                "json",
            ],
            capture_output=True,
            text=True,
            timeout=30,
        )

        # Verify workflow success
        assert result.returncode == 0
        assert output_file.exists()

        # Verify JSON is valid and contains expected content
        with open(output_file, encoding="utf-8") as f:
            data = json.load(f)
            content_str = json.dumps(data).lower()
            assert "audit" in content_str or "summary" in content_str

    @pytest.mark.slow
    def test_full_workflow_batch_processing(self, tmp_path):
        """Simulate real user workflow: batch process directory."""
        from docx import Document

        # User has directory of documents
        input_dir = tmp_path / "audit-documents"
        input_dir.mkdir()

        documents = [
            ("risk-assessment.docx", "Risk Assessment"),
            ("control-testing.docx", "Control Testing Results"),
            ("final-report.docx", "Final Audit Report"),
        ]

        for filename, heading in documents:
            file_path = input_dir / filename
            doc = Document()
            doc.add_heading(heading, level=1)
            doc.add_paragraph(f"Content for {heading}")
            doc.save(file_path)

        # User batch processes all documents
        output_dir = tmp_path / "extracted-data"

        result = subprocess.run(
            [
                "data-extract",
                "batch",
                str(input_dir),
                "--output",
                str(output_dir),
                "--format",
                "json",
                "--workers",
                "2",
            ],
            capture_output=True,
            text=True,
            timeout=60,
        )

        # Verify batch success
        assert result.returncode == 0
        assert output_dir.exists()

        # Verify all files processed
        output_files = list(output_dir.glob("*.json"))
        assert len(output_files) >= 3

        # Verify summary shown
        assert "Summary" in result.stdout or "successful" in result.stdout.lower()
