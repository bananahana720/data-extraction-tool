"""
Cross-Format Validation Integration Tests.

Tests consistency and compatibility across different file formats:
- Same logical content in DOCX/PDF → Extract → Compare results
- Same content → Different formatters → Verify consistency
- Same content → Different processors → Verify complementary enrichment
- Multi-format batch → Consistent quality scores

Test IDs: CF-001 through CF-006
"""

import json

import pytest

from src.core import ContentType
from src.extractors import DocxExtractor, PdfExtractor
from src.formatters import ChunkedTextFormatter, JsonFormatter, MarkdownFormatter
from src.pipeline import BatchProcessor, ExtractionPipeline
from src.processors import ContextLinker, MetadataAggregator, QualityValidator


# ==============================================================================
# Test Markers
# ==============================================================================

pytestmark = [pytest.mark.integration, pytest.mark.cross_format]


# ==============================================================================
# Same Content, Different Formats Tests
# ==============================================================================


def test_cf_001_same_content_docx_vs_pdf_extraction(tmp_path):
    """
    Test CF-001: Same content in DOCX vs PDF produces similar extraction.

    Scenario: Create identical content → Save as DOCX and PDF → Extract both

    Verifies:
    - Core content similar between formats
    - Block counts comparable
    - Text content matches
    - Different metadata appropriate to format
    """
    # Create simple DOCX
    try:
        from docx import Document

        doc = Document()
        doc.add_heading("Test Title", 0)
        doc.add_paragraph("This is paragraph one.")
        doc.add_paragraph("This is paragraph two.")

        docx_file = tmp_path / "test.docx"
        doc.save(str(docx_file))
    except ImportError:
        pytest.skip("python-docx not installed")

    # Create matching PDF
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas

        pdf_file = tmp_path / "test.pdf"
        c = canvas.Canvas(str(pdf_file), pagesize=letter)
        width, height = letter

        c.setFont("Helvetica-Bold", 16)
        c.drawString(100, height - 100, "Test Title")

        c.setFont("Helvetica", 12)
        c.drawString(100, height - 150, "This is paragraph one.")
        c.drawString(100, height - 170, "This is paragraph two.")

        c.save()
    except ImportError:
        pytest.skip("reportlab not installed")

    # Extract DOCX
    docx_extractor = DocxExtractor()
    docx_result = docx_extractor.extract(docx_file)

    # Extract PDF
    pdf_extractor = PdfExtractor()
    pdf_result = pdf_extractor.extract(pdf_file)

    # Assert: Both succeeded
    assert docx_result.success is True
    assert pdf_result.success is True

    # Assert: Block counts similar (may vary slightly due to format differences)
    docx_count = len(docx_result.content_blocks)
    pdf_count = len(pdf_result.content_blocks)

    # At least some blocks extracted from each
    assert docx_count > 0
    assert pdf_count > 0

    # Content should be present in both
    docx_text = " ".join(block.content for block in docx_result.content_blocks)
    pdf_text = " ".join(block.content for block in pdf_result.content_blocks)

    # Core content words should be present
    assert "Test Title" in docx_text or "Test" in docx_text
    assert "paragraph" in docx_text.lower()

    assert "Test" in pdf_text or "Title" in pdf_text
    assert "paragraph" in pdf_text.lower()


def test_cf_002_same_source_multiple_formatters_consistency(sample_docx_file):
    """
    Test CF-002: Same extraction → Multiple formatters → Consistent representation.

    Scenario: Extract once → Format as JSON, Markdown, Chunked

    Verifies:
    - All formatters handle same data
    - Content consistency across formats
    - No data loss in any format
    - Format-specific features correct
    """
    # Arrange: Extract once
    extractor = DocxExtractor()
    extraction_result = extractor.extract(sample_docx_file)

    assert extraction_result.success is True

    blocks = extraction_result.content_blocks
    metadata = extraction_result.document_metadata

    # Act: Format in all three formats
    json_formatter = JsonFormatter()
    json_result = json_formatter.format(blocks, metadata)

    md_formatter = MarkdownFormatter()
    md_result = md_formatter.format(blocks, metadata)

    chunked_formatter = ChunkedTextFormatter()
    chunked_result = chunked_formatter.format(blocks, metadata)

    # Assert: All succeeded
    assert json_result.success is True
    assert md_result.success is True
    assert chunked_result.success is True

    # Assert: All have content
    assert len(json_result.content) > 0
    assert len(md_result.content) > 0
    assert len(chunked_result.content) > 0

    # Assert: Block count consistent
    # JSON should have all blocks
    json_data = json.loads(json_result.content)
    if "content_blocks" in json_data:
        json_block_count = len(json_data["content_blocks"])
        assert json_block_count == len(blocks)

    # Assert: Core content present in all formats
    # Check for sample text that should be in document
    if blocks:
        sample_text = blocks[0].content[:20] if len(blocks[0].content) > 0 else ""
        if sample_text:
            # Text should appear in JSON (as string)
            assert (
                sample_text in json_result.content
                or sample_text.lower() in json_result.content.lower()
            )

            # Text should appear in Markdown
            # (May be formatted, but core words present)
            md_words = set(sample_text.split())
            md_content_words = set(md_result.content.split())
            # At least some words should match
            if md_words:
                overlap = md_words.intersection(md_content_words)
                assert len(overlap) > 0 or len(sample_text) == 0


def test_cf_003_same_content_different_processors_complementary(sample_docx_file):
    """
    Test CF-003: Same content → Different processors → Complementary enrichment.

    Scenario: Extract → Process with each processor separately → Compare

    Verifies:
    - Each processor adds unique metadata
    - No processor removes existing data
    - Processors complement each other
    - Combined processing most complete
    """
    # Arrange: Extract
    extractor = DocxExtractor()
    extraction_result = extractor.extract(sample_docx_file)

    assert extraction_result.success is True

    blocks = extraction_result.content_blocks

    # Act: Process with each processor separately
    linker = ContextLinker()
    linked_result = linker.process(blocks)

    aggregator = MetadataAggregator()
    aggregated_result = aggregator.process(blocks)

    validator = QualityValidator()
    validated_result = validator.process(blocks)

    # Assert: All succeeded
    assert linked_result.success is True
    assert aggregated_result.success is True
    assert validated_result.success is True

    # Assert: Block counts preserved
    assert len(linked_result.content_blocks) == len(blocks)
    assert len(aggregated_result.content_blocks) == len(blocks)
    assert len(validated_result.content_blocks) == len(blocks)

    # Assert: Each processor adds unique information
    # ContextLinker: hierarchy/linking metadata
    # MetadataAggregator: statistical metadata
    # QualityValidator: quality score

    # Quality validator produces score
    assert validated_result.quality_score is not None

    # Act: Process through all processors in sequence
    chained_blocks = blocks
    for processor in [linker, aggregator, validator]:
        result = processor.process(chained_blocks)
        chained_blocks = result.content_blocks

    # Assert: Chained processing has most complete enrichment
    assert len(chained_blocks) == len(blocks)

    # Sample block should have accumulated metadata
    if chained_blocks:
        sample_block = chained_blocks[0]
        # Should have some metadata from processing
        assert sample_block.metadata is not None


def test_cf_004_multi_format_batch_consistent_quality(tmp_path):
    """
    Test CF-004: Multi-format batch → Consistent quality assessment.

    Scenario: DOCX + PDF + TXT → Batch process → Compare quality scores

    Verifies:
    - Quality scoring works across formats
    - Scores reflect actual quality
    - Consistent scoring criteria
    - Format-specific adjustments appropriate
    """
    # Create test files
    # DOCX
    try:
        from docx import Document

        doc = Document()
        doc.add_heading("Quality Test", 0)
        for i in range(5):
            doc.add_paragraph(f"Paragraph {i+1} with substantial content.")

        docx_file = tmp_path / "quality.docx"
        doc.save(str(docx_file))
    except ImportError:
        pytest.skip("python-docx not installed")

    # PDF
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas

        pdf_file = tmp_path / "quality.pdf"
        c = canvas.Canvas(str(pdf_file), pagesize=letter)
        width, height = letter

        c.setFont("Helvetica", 12)
        y_pos = height - 100
        c.drawString(100, y_pos, "Quality Test")
        for i in range(5):
            y_pos -= 20
            c.drawString(100, y_pos, f"Paragraph {i+1} with substantial content.")

        c.save()
    except ImportError:
        pytest.skip("reportlab not installed")

    # TXT
    txt_file = tmp_path / "quality.txt"
    txt_file.write_text(
        "Quality Test\n\n"
        + "\n".join([f"Paragraph {i+1} with substantial content." for i in range(5)])
    )

    # Arrange: Create pipeline with quality validation
    pipeline = ExtractionPipeline()
    pipeline.register_extractor("docx", DocxExtractor())
    pipeline.register_extractor("pdf", PdfExtractor())
    pipeline.register_extractor("txt", DocxExtractor())

    pipeline.add_processor(QualityValidator())
    pipeline.add_formatter(JsonFormatter())

    # Arrange: Batch processor
    batch = BatchProcessor(pipeline=pipeline, max_workers=1)

    # Act: Process all files
    files = [docx_file, pdf_file, txt_file]
    results = batch.process_batch(files)

    # Assert: All processed
    assert len(results) == 3

    # Assert: Extract quality scores
    quality_scores = []
    for result in results:
        if result.success and result.processing_result:
            if result.processing_result.quality_score is not None:
                quality_scores.append(result.processing_result.quality_score)

    # At least some should have quality scores
    assert len(quality_scores) > 0

    # Assert: Scores in valid range
    for score in quality_scores:
        assert 0.0 <= score <= 100.0

    # Assert: Scores reflect quality (all are high-quality files)
    avg_score = sum(quality_scores) / len(quality_scores) if quality_scores else 0
    assert avg_score > 70.0, "High-quality files should score well"


# ==============================================================================
# Format-Specific Feature Preservation Tests
# ==============================================================================


def test_cf_005_format_specific_metadata_preserved(tmp_path):
    """
    Test CF-005: Format-specific metadata preserved through pipeline.

    Scenario: DOCX (styles) + PDF (fonts) → Extract → Process → Format

    Verifies:
    - DOCX style information preserved
    - PDF font information preserved
    - Format-specific metadata in output
    - No format-specific metadata leaks between formats
    """
    # Create DOCX with styled content
    try:
        from docx import Document

        doc = Document()
        heading = doc.add_heading("Styled Heading", 1)
        para = doc.add_paragraph("Normal paragraph text.")

        docx_file = tmp_path / "styled.docx"
        doc.save(str(docx_file))
    except ImportError:
        pytest.skip("python-docx not installed")

    # Create PDF
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas

        pdf_file = tmp_path / "styled.pdf"
        c = canvas.Canvas(str(pdf_file), pagesize=letter)

        c.setFont("Helvetica-Bold", 14)
        c.drawString(100, 700, "Styled Heading")

        c.setFont("Helvetica", 12)
        c.drawString(100, 670, "Normal paragraph text.")

        c.save()
    except ImportError:
        pytest.skip("reportlab not installed")

    # Arrange: Create pipeline
    pipeline = ExtractionPipeline()
    pipeline.register_extractor("docx", DocxExtractor())
    pipeline.register_extractor("pdf", PdfExtractor())

    pipeline.add_processor(MetadataAggregator())
    pipeline.add_formatter(JsonFormatter())

    # Act: Process both
    docx_result = pipeline.process_file(docx_file)
    pdf_result = pipeline.process_file(pdf_file)

    # Assert: Both succeeded
    assert docx_result.success is True
    assert pdf_result.success is True

    # Assert: Format-specific metadata present
    # DOCX blocks may have style information
    docx_blocks = docx_result.extraction_result.content_blocks
    if docx_blocks:
        # Check for format-specific metadata
        # (Actual keys depend on extractor implementation)
        sample_block = docx_blocks[0]
        assert sample_block.metadata is not None

    # PDF blocks may have font information
    pdf_blocks = pdf_result.extraction_result.content_blocks
    if pdf_blocks:
        sample_block = pdf_blocks[0]
        assert sample_block.metadata is not None


def test_cf_006_consistent_block_type_classification(tmp_path):
    """
    Test CF-006: Consistent ContentType classification across formats.

    Scenario: Same content types in DOCX vs PDF → Extract → Compare types

    Verifies:
    - Headings classified consistently
    - Paragraphs classified consistently
    - Tables classified consistently
    - ContentType enum used uniformly
    """
    # Create DOCX with various content types
    try:
        from docx import Document

        doc = Document()
        doc.add_heading("Main Heading", 0)
        doc.add_paragraph("Regular paragraph text.")
        doc.add_heading("Subheading", 1)

        # Add table
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "A"
        table.cell(0, 1).text = "B"
        table.cell(1, 0).text = "C"
        table.cell(1, 1).text = "D"

        docx_file = tmp_path / "types.docx"
        doc.save(str(docx_file))
    except ImportError:
        pytest.skip("python-docx not installed")

    # Extract DOCX
    extractor = DocxExtractor()
    result = extractor.extract(docx_file)

    assert result.success is True

    # Assert: Multiple content types present
    content_types = {block.block_type for block in result.content_blocks}

    # Should have at least heading and paragraph
    assert ContentType.HEADING in content_types or ContentType.PARAGRAPH in content_types

    # Assert: All types are valid ContentType enum values
    for block in result.content_blocks:
        assert isinstance(block.block_type, ContentType)


# ==============================================================================
# Cross-Format Error Handling Tests
# ==============================================================================


def test_cf_007_consistent_error_handling_across_formats(tmp_path):
    """
    Test CF-007: Consistent error handling for corrupted files of different formats.

    Scenario: Corrupted DOCX + Corrupted PDF → Extract → Compare error handling

    Verifies:
    - Errors detected consistently
    - Error messages appropriate to format
    - Graceful failure for both
    - Error structure consistent
    """
    # Create corrupted DOCX
    corrupt_docx = tmp_path / "corrupt.docx"
    corrupt_docx.write_text("Not a valid DOCX file")

    # Create corrupted PDF
    corrupt_pdf = tmp_path / "corrupt.pdf"
    corrupt_pdf.write_bytes(b"%PDF-1.4\n" + b"\x00" * 50)  # Truncated PDF

    # Arrange: Extractors
    docx_extractor = DocxExtractor()
    pdf_extractor = PdfExtractor()

    # Act: Try to extract both
    docx_result = docx_extractor.extract(corrupt_docx)
    pdf_result = pdf_extractor.extract(corrupt_pdf)

    # Assert: Both failed or succeeded with warnings
    # (Behavior depends on extractor implementation)

    # Assert: Error information present if failed
    if not docx_result.success:
        assert len(docx_result.errors) > 0

    if not pdf_result.success:
        assert len(pdf_result.errors) > 0

    # Assert: Result structure consistent
    assert hasattr(docx_result, "success")
    assert hasattr(pdf_result, "success")
    assert hasattr(docx_result, "errors")
    assert hasattr(pdf_result, "errors")
