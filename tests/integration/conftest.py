"""
Pytest fixtures for integration tests.

This module provides shared fixtures for integration testing including:
- Sample document files (DOCX, PDF, PPTX, XLSX)
- Corrupted files for error testing
- Configured pipelines
- CLI runners
- Temporary directories
"""

import shutil
import tempfile
from pathlib import Path
from typing import Generator

import pytest


# ==============================================================================
# Sample Document Fixtures
# ==============================================================================


@pytest.fixture
def sample_docx_file(tmp_path: Path) -> Path:
    """
    Create a sample DOCX file for testing.

    Generates a valid Word document with:
    - Headings (multiple levels)
    - Paragraphs with text
    - Table
    - Document metadata

    Args:
        tmp_path: Pytest built-in temporary directory

    Returns:
        Path to generated DOCX file
    """
    try:
        from docx import Document
    except ImportError as e:
        pytest.skip(
            f"python-docx not installed - install with: pip install python-docx\n"
            f"This package is required for DOCX test fixtures. Error: {e}"
        )

    doc = Document()

    # Add metadata
    core_props = doc.core_properties
    core_props.title = "Test Document"
    core_props.author = "Test Author"
    core_props.subject = "Integration Testing"

    # Add content
    doc.add_heading("Test Document", 0)
    doc.add_paragraph(
        "This is a test document for integration testing. "
        "It contains multiple types of content blocks."
    )

    doc.add_heading("Section 1: Introduction", 1)
    doc.add_paragraph(
        "This section introduces the test document. "
        "It has several paragraphs to test extraction."
    )
    doc.add_paragraph("Second paragraph in the introduction section.")

    doc.add_heading("Section 2: Data Table", 1)
    doc.add_paragraph("Below is a sample table:")

    # Add table
    table = doc.add_table(rows=3, cols=3)
    table.style = "Table Grid"

    # Header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Column 1"
    hdr_cells[1].text = "Column 2"
    hdr_cells[2].text = "Column 3"

    # Data rows
    row1 = table.rows[1].cells
    row1[0].text = "A1"
    row1[1].text = "B1"
    row1[2].text = "C1"

    row2 = table.rows[2].cells
    row2[0].text = "A2"
    row2[1].text = "B2"
    row2[2].text = "C2"

    doc.add_heading("Section 3: Conclusion", 1)
    doc.add_paragraph("This concludes the test document.")

    # Save file
    file_path = tmp_path / "test_document.docx"
    # python-docx 1.2.0+ supports Path objects natively
    doc.save(file_path)

    return file_path


@pytest.fixture
def sample_pdf_file(tmp_path: Path) -> Path:
    """
    Create a sample PDF file for testing.

    Generates a simple PDF with text content.

    Args:
        tmp_path: Pytest built-in temporary directory

    Returns:
        Path to generated PDF file
    """
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas
    except ImportError as e:
        pytest.skip(
            f"reportlab not installed - install with: pip install reportlab\n"
            f"This package is required for PDF test fixtures. Error: {e}"
        )

    file_path = tmp_path / "test_document.pdf"

    # Create PDF
    # reportlab 3.5+ supports Path objects natively
    c = canvas.Canvas(file_path, pagesize=letter)
    width, height = letter

    # Add content
    c.setFont("Helvetica-Bold", 16)
    c.drawString(100, height - 100, "Test PDF Document")

    c.setFont("Helvetica", 12)
    c.drawString(100, height - 150, "This is a test PDF for integration testing.")
    c.drawString(100, height - 170, "It contains multiple lines of text.")

    c.drawString(100, height - 220, "Section 1: Introduction")
    c.drawString(100, height - 240, "This is the introduction section.")

    c.drawString(100, height - 290, "Section 2: Content")
    c.drawString(100, height - 310, "This section contains the main content.")
    c.drawString(100, height - 330, "Multiple paragraphs are included.")

    # Add second page
    c.showPage()
    c.setFont("Helvetica", 12)
    c.drawString(100, height - 100, "Page 2")
    c.drawString(100, height - 120, "This is the second page of the document.")

    c.save()

    return file_path


@pytest.fixture
def sample_text_file(tmp_path: Path) -> Path:
    """
    Create a simple text file for testing.

    Args:
        tmp_path: Pytest built-in temporary directory

    Returns:
        Path to generated text file
    """
    file_path = tmp_path / "test_document.txt"
    file_path.write_text(
        "Test Document\n\n"
        "This is a test document for integration testing.\n\n"
        "Section 1: Introduction\n"
        "This section introduces the test document.\n\n"
        "Section 2: Content\n"
        "This is the main content section.\n"
        "It has multiple paragraphs.\n\n"
        "Section 3: Conclusion\n"
        "This concludes the test document.\n"
    )
    return file_path


@pytest.fixture
def large_docx_file(tmp_path: Path) -> Path:
    """
    Create a large DOCX file for performance testing.

    Generates a document with 100+ paragraphs (~1MB).

    Args:
        tmp_path: Pytest built-in temporary directory

    Returns:
        Path to generated large DOCX file
    """
    try:
        from docx import Document
    except ImportError as e:
        pytest.skip(
            f"python-docx not installed - install with: pip install python-docx\n"
            f"This package is required for DOCX test fixtures. Error: {e}"
        )

    doc = Document()
    doc.add_heading("Large Test Document", 0)

    # Add 100 paragraphs
    for i in range(100):
        doc.add_heading(f"Section {i+1}", 1)
        doc.add_paragraph(f"This is paragraph {i+1} of the large test document. " * 10)

    file_path = tmp_path / "large_document.docx"
    # python-docx 1.2.0+ supports Path objects natively
    doc.save(file_path)

    return file_path


# ==============================================================================
# Corrupted File Fixtures
# ==============================================================================


@pytest.fixture
def corrupted_docx_file(tmp_path: Path) -> Path:
    """
    Create a corrupted DOCX file for error testing.

    Creates a file with .docx extension but invalid content.

    Args:
        tmp_path: Pytest built-in temporary directory

    Returns:
        Path to corrupted DOCX file
    """
    file_path = tmp_path / "corrupted.docx"
    file_path.write_text("This is not a valid DOCX file")
    return file_path


@pytest.fixture
def corrupted_pdf_file(tmp_path: Path) -> Path:
    """
    Create a corrupted PDF file for error testing.

    Creates a file with .pdf extension but invalid content.

    Args:
        tmp_path: Pytest built-in temporary directory

    Returns:
        Path to corrupted PDF file
    """
    file_path = tmp_path / "corrupted.pdf"
    file_path.write_bytes(b"%PDF-1.4\n" + b"\x00" * 100)  # Truncated PDF
    return file_path


@pytest.fixture
def empty_docx_file(tmp_path: Path) -> Path:
    """
    Create an empty DOCX file for edge case testing.

    Args:
        tmp_path: Pytest built-in temporary directory

    Returns:
        Path to empty DOCX file
    """
    try:
        from docx import Document
    except ImportError as e:
        pytest.skip(
            f"python-docx not installed - install with: pip install python-docx\n"
            f"This package is required for DOCX test fixtures. Error: {e}"
        )

    doc = Document()
    file_path = tmp_path / "empty.docx"
    # python-docx 1.2.0+ supports Path objects natively
    doc.save(file_path)

    return file_path


# ==============================================================================
# Multiple File Fixtures
# ==============================================================================


@pytest.fixture
def multiple_test_files(tmp_path: Path, sample_docx_file, sample_text_file) -> list[Path]:
    """
    Create multiple test files for batch processing.

    Args:
        tmp_path: Pytest built-in temporary directory
        sample_docx_file: Sample DOCX file fixture
        sample_text_file: Sample text file fixture

    Returns:
        List of Paths to test files
    """
    files = []

    # Copy sample files to batch directory
    batch_dir = tmp_path / "batch"
    batch_dir.mkdir()

    # Create 5 test files
    for i in range(1, 6):
        if i % 2 == 0:
            # Even numbers: copy DOCX
            dest = batch_dir / f"file_{i}.docx"
            shutil.copy(sample_docx_file, dest)
        else:
            # Odd numbers: copy text
            dest = batch_dir / f"file_{i}.txt"
            shutil.copy(sample_text_file, dest)

        files.append(dest)

    return files


@pytest.fixture
def batch_test_directory(multiple_test_files) -> Path:
    """
    Get the directory containing multiple test files.

    Args:
        multiple_test_files: List of test files

    Returns:
        Path to directory containing files
    """
    return multiple_test_files[0].parent


# ==============================================================================
# Pipeline Fixtures
# ==============================================================================


@pytest.fixture
def configured_pipeline():
    """
    Create a fully configured ExtractionPipeline for testing.

    Returns:
        Configured ExtractionPipeline instance
    """
    from src.pipeline import ExtractionPipeline
    from src.extractors import DocxExtractor, PdfExtractor, TextFileExtractor
    from src.processors import ContextLinker, MetadataAggregator, QualityValidator
    from src.formatters import JsonFormatter, MarkdownFormatter, ChunkedTextFormatter

    pipeline = ExtractionPipeline()

    # Register extractors
    pipeline.register_extractor("docx", DocxExtractor())
    pipeline.register_extractor("pdf", PdfExtractor())
    pipeline.register_extractor("txt", TextFileExtractor())  # Use correct text extractor

    # Add processors
    pipeline.add_processor(ContextLinker())
    pipeline.add_processor(MetadataAggregator())
    pipeline.add_processor(QualityValidator())

    # Add formatters
    pipeline.add_formatter(JsonFormatter())
    pipeline.add_formatter(MarkdownFormatter())
    pipeline.add_formatter(ChunkedTextFormatter())

    return pipeline


@pytest.fixture
def batch_processor(configured_pipeline):
    """
    Create a configured BatchProcessor for testing.

    Args:
        configured_pipeline: Configured pipeline fixture

    Returns:
        Configured BatchProcessor instance
    """
    from src.pipeline import BatchProcessor

    return BatchProcessor(
        pipeline=configured_pipeline,
        max_workers=2,  # Use 2 workers for testing
    )


# ==============================================================================
# CLI Fixtures
# ==============================================================================


@pytest.fixture
def cli_runner():
    """
    Create a Click CLI test runner.

    Returns:
        CliRunner instance for testing CLI commands
    """
    from click.testing import CliRunner

    return CliRunner()


@pytest.fixture
def config_file(tmp_path: Path) -> Path:
    """
    Create a sample configuration file for testing.

    Args:
        tmp_path: Pytest built-in temporary directory

    Returns:
        Path to config file
    """
    config_content = """
pipeline:
  max_workers: 4
  timeout_per_file: 300

extractors:
  docx:
    skip_empty: true
  pdf:
    ocr_enabled: false

processors:
  - ContextLinker
  - MetadataAggregator
  - QualityValidator

formatters:
  - type: json
    pretty_print: true
  - type: markdown
    include_metadata: true
"""

    config_path = tmp_path / "config.yaml"
    config_path.write_text(config_content)

    return config_path


@pytest.fixture
def invalid_config_file(tmp_path: Path) -> Path:
    """
    Create an invalid configuration file for error testing.

    Args:
        tmp_path: Pytest built-in temporary directory

    Returns:
        Path to invalid config file
    """
    invalid_content = """
pipeline:
  max_workers: "not a number"  # Invalid type
  timeout_per_file: -100  # Invalid value
"""

    config_path = tmp_path / "invalid_config.yaml"
    config_path.write_text(invalid_content)

    return config_path


# ==============================================================================
# Output Directory Fixtures
# ==============================================================================


@pytest.fixture
def output_directory(tmp_path: Path) -> Path:
    """
    Create a temporary output directory for test results.

    Args:
        tmp_path: Pytest built-in temporary directory

    Returns:
        Path to output directory
    """
    output_dir = tmp_path / "output"
    output_dir.mkdir()
    return output_dir


# ==============================================================================
# Progress Tracking Fixtures
# ==============================================================================


@pytest.fixture
def progress_tracker():
    """
    Create a list to collect progress updates during tests.

    Returns:
        Tuple of (progress_list, callback_function)
    """
    progress_updates = []

    def callback(status: dict) -> None:
        """Record progress updates."""
        progress_updates.append(status.copy())

    return progress_updates, callback


# ==============================================================================
# Performance Testing Fixtures
# ==============================================================================


@pytest.fixture
def performance_timer():
    """
    Create a simple timer for performance testing.

    Returns:
        Timer object with start/stop methods
    """
    import time

    class Timer:
        def __init__(self):
            self.start_time = None
            self.end_time = None

        def start(self):
            self.start_time = time.perf_counter()

        def stop(self):
            self.end_time = time.perf_counter()

        @property
        def elapsed(self) -> float:
            if self.start_time is None or self.end_time is None:
                return 0.0
            return self.end_time - self.start_time

    return Timer()


# ==============================================================================
# Cleanup Fixtures
# ==============================================================================


@pytest.fixture(autouse=True)
def cleanup_temp_files(tmp_path):
    """
    Automatically cleanup temporary files after each test.

    This fixture runs automatically for all tests in this package.
    Tracks files created outside of tmp_path and cleans them up.

    Files created in tmp_path are cleaned up automatically by pytest.
    This fixture handles edge cases where tests might create files elsewhere.
    """
    import glob
    import os

    # Track current directory and any potential test artifacts
    original_cwd = os.getcwd()
    test_artifacts = []

    # Common test artifact patterns to clean up
    artifact_patterns = [
        "empty.txt",  # Created by edge case tests
        "minimal.txt",  # Created by edge case tests
        "corrupted.docx",  # Created by edge case tests
        "test_*.tmp",  # Any test temp files
    ]

    yield

    # Cleanup: Remove any test artifacts created in the current directory
    try:
        for pattern in artifact_patterns:
            for file_path in glob.glob(pattern):
                try:
                    if os.path.isfile(file_path):
                        os.unlink(file_path)
                except Exception:
                    pass  # Ignore cleanup errors - don't fail tests due to cleanup

        # Ensure we're back in the original directory
        os.chdir(original_cwd)
    except Exception:
        pass  # Cleanup is best-effort; don't fail tests if it fails
