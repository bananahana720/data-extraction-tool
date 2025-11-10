"""
Test script for DocxExtractor.

Creates a sample DOCX file and tests extraction.
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE

import sys

sys.path.insert(0, str(Path(__file__).parent / "src"))

from extractors import DocxExtractor


def create_test_docx(file_path: Path) -> None:
    """Create a test DOCX file with various content types."""
    doc = Document()

    # Title
    title = doc.add_heading("Test Document", level=1)

    # Introduction paragraph
    intro = doc.add_paragraph(
        "This is a test document created to demonstrate the DocxExtractor. "
        "It contains multiple paragraphs with different content types."
    )

    # Section heading
    doc.add_heading("Section 1: Basic Content", level=2)

    # Regular paragraphs
    doc.add_paragraph(
        "This is a regular paragraph with normal text. "
        "It should be extracted as ContentType.PARAGRAPH."
    )

    doc.add_paragraph(
        "Here is another paragraph. This one has multiple sentences. "
        "Each sentence adds to the word count."
    )

    # Another section
    doc.add_heading("Section 2: More Content", level=2)

    # More content
    doc.add_paragraph(
        "The extractor should maintain the sequence order of all paragraphs. "
        "Position information helps reconstruct document structure."
    )

    # Empty paragraph (should be skipped by default)
    doc.add_paragraph("")

    # Final paragraph
    doc.add_paragraph(
        "This is the final paragraph in our test document. "
        "It demonstrates that extraction works end-to-end."
    )

    # Save document
    doc.save(file_path)
    print(f"Created test document: {file_path}")


def test_docx_extractor():
    """Test the DocxExtractor with a sample file."""
    print("=" * 70)
    print("DOCX EXTRACTOR TEST")
    print("=" * 70)

    # Create test file
    test_file = Path("test_sample.docx")
    create_test_docx(test_file)

    # Create extractor
    extractor = DocxExtractor()

    # Check format support
    print(f"\n[CHECK] Format supported: {extractor.supports_format(test_file)}")
    print(f"[CHECK] Format name: {extractor.get_format_name()}")
    print(f"[CHECK] Supported extensions: {extractor.get_supported_extensions()}")

    # Extract content
    print(f"\n[EXTRACT] Extracting content from {test_file}...")
    result = extractor.extract(test_file)

    # Display results
    print("\n" + "=" * 70)
    if result.success:
        print("[SUCCESS] Extraction successful!")
        print("=" * 70)

        # Document metadata
        meta = result.document_metadata
        print(f"\nDocument Metadata:")
        print(f"  File: {meta.source_file.name}")
        print(f"  Format: {meta.file_format}")
        print(f"  Size: {meta.file_size_bytes:,} bytes")
        print(f"  Word count: {meta.word_count}")
        print(f"  Character count: {meta.character_count}")
        print(f"  Hash: {meta.file_hash[:16]}...")

        if meta.title:
            print(f"  Title: {meta.title}")
        if meta.author:
            print(f"  Author: {meta.author}")

        print(f"\nExtraction Statistics:")
        print(f"  Content blocks: {len(result.content_blocks)}")
        print(f"  Images: {len(result.images)}")
        print(f"  Tables: {len(result.tables)}")

        if result.warnings:
            print(f"\n[WARNINGS] {len(result.warnings)} warning(s):")
            for warning in result.warnings:
                print(f"  - {warning}")

        # Show content blocks
        print(f"\n{'=' * 70}")
        print("Content Blocks:")
        print("=" * 70)

        for idx, block in enumerate(result.content_blocks):
            print(f"\n[Block {idx + 1}]")
            print(f"  Type: {block.block_type.value}")
            print(f"  Position: {block.position}")
            print(f"  Confidence: {block.confidence}")

            # Show metadata
            if block.metadata:
                print(f"  Metadata:")
                for key, value in block.metadata.items():
                    print(f"    {key}: {value}")

            # Show content preview
            content_preview = block.content[:100]
            if len(block.content) > 100:
                content_preview += "..."
            print(f"  Content: {content_preview}")

        print("\n" + "=" * 70)
        print("[TEST PASSED]")
        print("=" * 70)

    else:
        print("[FAILED] Extraction failed:")
        print("=" * 70)
        for error in result.errors:
            print(f"  - {error}")

        if result.warnings:
            print(f"\nWarnings:")
            for warning in result.warnings:
                print(f"  - {warning}")

        print("\n" + "=" * 70)
        print("[TEST FAILED]")
        print("=" * 70)

    # Cleanup
    print(f"\n[CLEANUP] Removing test file...")
    test_file.unlink()
    print("[COMPLETE]")


if __name__ == "__main__":
    test_docx_extractor()
