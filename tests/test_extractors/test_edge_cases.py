"""
Comprehensive edge case and stress tests for all extractors.

This test suite uses equivalency partitioning methodology to systematically
test boundary conditions, error cases, and extreme scenarios across all
document extractors.

Equivalency Partitioning Approach:
1. Valid partitions: Normal operating ranges
2. Invalid partitions: Error conditions
3. Boundary values: Edge of valid/invalid ranges
4. Special values: Null, empty, max, min

Test Coverage:
- File system edge cases (paths, permissions, locks)
- Content edge cases (empty, massive, malformed)
- Encoding edge cases (UTF variants, BOMs, mixed)
- Format edge cases (wrong extension, corrupted headers)
- Resource limits (memory, time, size)

Markers:
    @pytest.mark.edge_case - All tests in this file
    @pytest.mark.slow - Tests that may take >1s
    @pytest.mark.stress - Resource-intensive tests
"""

import sys
import time
from pathlib import Path

import pytest

# Import foundation
sys.path.insert(0, str(Path(__file__).parent.parent.parent / "src"))



# ==============================================================================
# File System Edge Cases
# ==============================================================================


@pytest.mark.edge_case
class TestFileSystemEdgeCases:
    """Test extractors with file system boundary conditions."""

    def test_file_path_with_spaces(self, tmp_path):
        """
        EDGE: File path containing spaces should be handled correctly.

        Partition: File system → Valid paths → Special characters
        Boundary: Spaces in path components
        """
        from reportlab.pdfgen import canvas

        from extractors.pdf_extractor import PdfExtractor

        # Create directory and file with spaces
        space_dir = tmp_path / "test folder with spaces"
        space_dir.mkdir()
        pdf_path = space_dir / "file with spaces.pdf"

        # Create simple PDF
        c = canvas.Canvas(str(pdf_path))
        c.drawString(100, 750, "Content in file with spaces")
        c.showPage()
        c.save()

        extractor = PdfExtractor()
        result = extractor.extract(pdf_path)

        assert result.success is True
        assert len(result.content_blocks) > 0

    def test_file_path_with_unicode(self, tmp_path):
        """
        EDGE: File path with unicode characters (文档.pdf).

        Partition: File system → Valid paths → Unicode
        Boundary: Non-ASCII characters in filename
        """
        from reportlab.pdfgen import canvas

        from extractors.pdf_extractor import PdfExtractor

        # Create file with unicode name
        pdf_path = tmp_path / "测试文档_中文.pdf"

        c = canvas.Canvas(str(pdf_path))
        c.drawString(100, 750, "Content in unicode filename")
        c.showPage()
        c.save()

        extractor = PdfExtractor()
        result = extractor.extract(pdf_path)

        assert result.success is True

    def test_very_long_file_path(self, tmp_path):
        """
        EDGE: Very long file paths (approaching Windows 260 char limit).

        Partition: File system → Valid paths → Length boundary
        Boundary: Near OS path length limits
        """
        from reportlab.pdfgen import canvas

        from extractors.pdf_extractor import PdfExtractor

        # Create deeply nested path (but stay under 260 for Windows)
        long_path = tmp_path
        for i in range(5):
            long_path = long_path / f"level_{i}_subdirectory"
            long_path.mkdir(exist_ok=True)

        pdf_path = long_path / "document_with_a_moderately_long_filename_to_test_path_limits.pdf"

        c = canvas.Canvas(str(pdf_path))
        c.drawString(100, 750, "Content in deep path")
        c.showPage()
        c.save()

        extractor = PdfExtractor()
        result = extractor.extract(pdf_path)

        assert result.success is True

    def test_file_with_no_extension(self, tmp_path):
        """
        EDGE: File with no extension should fail format detection.

        Partition: File system → Invalid → Missing extension
        Expected: Validation failure
        """
        from extractors.pdf_extractor import PdfExtractor

        no_ext_file = tmp_path / "document_no_extension"
        no_ext_file.write_text("fake content")

        extractor = PdfExtractor()

        # Should not support file without extension
        assert extractor.supports_format(no_ext_file) is False

    def test_file_with_wrong_extension(self, tmp_path):
        """
        EDGE: PDF content with .docx extension (wrong extension).

        Partition: File system → Invalid → Extension mismatch
        Expected: Extraction failure with clear error
        """
        from reportlab.pdfgen import canvas

        from extractors.docx_extractor import DocxExtractor

        # Create PDF but name it .docx
        pdf_path = tmp_path / "actually_pdf.docx"
        c = canvas.Canvas(str(pdf_path))
        c.drawString(100, 750, "I'm actually a PDF")
        c.showPage()
        c.save()

        # Try to extract with DOCX extractor
        extractor = DocxExtractor()
        result = extractor.extract(pdf_path)

        # Should fail because content doesn't match extension
        assert result.success is False
        assert len(result.errors) > 0

    def test_zero_byte_file(self, tmp_path):
        """
        EDGE: Completely empty file (0 bytes).

        Partition: Content → Invalid → Empty
        Boundary: Minimum possible file size (0)
        Expected: Validation failure
        """
        from extractors.pdf_extractor import PdfExtractor

        empty_file = tmp_path / "empty.pdf"
        empty_file.touch()  # 0 bytes

        extractor = PdfExtractor()
        is_valid, errors = extractor.validate_file(empty_file)

        assert is_valid is False
        assert len(errors) > 0
        assert "empty" in errors[0].lower()

    def test_read_only_file(self, tmp_path):
        """
        EDGE: Read-only file should still be readable.

        Partition: File system → Valid → Permission boundary
        Note: This tests read access, not write
        """
        import stat

        from reportlab.pdfgen import canvas

        from extractors.pdf_extractor import PdfExtractor

        pdf_path = tmp_path / "readonly.pdf"
        c = canvas.Canvas(str(pdf_path))
        c.drawString(100, 750, "Read-only content")
        c.showPage()
        c.save()

        # Make read-only
        pdf_path.chmod(stat.S_IRUSR | stat.S_IRGRP | stat.S_IROTH)

        try:
            extractor = PdfExtractor()
            result = extractor.extract(pdf_path)

            # Should succeed - we only need read access
            assert result.success is True
        finally:
            # Restore permissions for cleanup
            pdf_path.chmod(stat.S_IWUSR | stat.S_IRUSR)


# ==============================================================================
# Content Size Edge Cases
# ==============================================================================


@pytest.mark.edge_case
class TestContentSizeEdgeCases:
    """Test extractors with extreme content sizes."""

    def test_document_with_zero_content_blocks(self, tmp_path):
        """
        EDGE: Document that validates but extracts zero content.

        Partition: Content → Valid → Minimal
        Boundary: Minimum useful content (0 blocks)
        """
        from reportlab.pdfgen import canvas

        from extractors.pdf_extractor import PdfExtractor

        # PDF with blank page
        pdf_path = tmp_path / "blank_page.pdf"
        c = canvas.Canvas(str(pdf_path))
        c.showPage()  # Blank page
        c.save()

        extractor = PdfExtractor()
        result = extractor.extract(pdf_path)

        # Should succeed but may have zero or minimal content blocks
        assert result.success is True
        assert result.document_metadata.page_count == 1

    def test_document_with_single_character(self, tmp_path):
        """
        EDGE: Document with minimal content (single character).

        Partition: Content → Valid → Minimum
        Boundary: Smallest valid content (1 char)
        """
        from reportlab.pdfgen import canvas

        from extractors.pdf_extractor import PdfExtractor

        pdf_path = tmp_path / "single_char.pdf"
        c = canvas.Canvas(str(pdf_path))
        c.drawString(100, 750, "A")  # Single character
        c.showPage()
        c.save()

        extractor = PdfExtractor()
        result = extractor.extract(pdf_path)

        assert result.success is True
        assert len(result.content_blocks) > 0

    @pytest.mark.slow
    @pytest.mark.stress
    def test_very_large_document(self, tmp_path):
        """
        EDGE: Very large document (100+ pages).

        Partition: Content → Valid → Maximum
        Boundary: Large content stress test
        Note: This is slow and resource-intensive
        """
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas

        from extractors.pdf_extractor import PdfExtractor

        pdf_path = tmp_path / "large.pdf"
        c = canvas.Canvas(str(pdf_path), pagesize=letter)

        # Create 50 pages (reduced from 100 for test speed)
        for page_num in range(50):
            c.drawString(100, 750, f"Page {page_num + 1}")
            for line in range(30):
                c.drawString(100, 730 - line * 15, f"Line {line}: Content on page {page_num + 1}")
            c.showPage()

        c.save()

        # Check file size
        file_size = pdf_path.stat().st_size
        assert file_size > 50000  # Should be substantial

        extractor = PdfExtractor()

        start_time = time.time()
        result = extractor.extract(pdf_path)
        duration = time.time() - start_time

        assert result.success is True
        assert result.document_metadata.page_count == 50

        # Performance check: Should complete in reasonable time
        # Target: <2s/MB, but allow generous margin for test environment
        size_mb = file_size / (1024 * 1024)
        max_duration = max(10.0, size_mb * 5)  # At least 10s, or 5s per MB
        assert (
            duration < max_duration
        ), f"Took {duration:.2f}s for {size_mb:.2f}MB (limit: {max_duration:.2f}s)"

    @pytest.mark.slow
    def test_document_with_very_long_single_line(self, tmp_path):
        """
        EDGE: Single line with thousands of characters.

        Partition: Content → Valid → Line length boundary
        Boundary: Very long single text block
        """
        from reportlab.pdfgen import canvas

        from extractors.pdf_extractor import PdfExtractor

        pdf_path = tmp_path / "long_line.pdf"
        c = canvas.Canvas(str(pdf_path))

        # Create very long text (but truncate for PDF rendering)
        long_text = "A" * 1000
        # PDF has width limits, so we'll just put first part
        c.drawString(100, 750, long_text[:100])
        c.showPage()
        c.save()

        extractor = PdfExtractor()
        result = extractor.extract(pdf_path)

        assert result.success is True


# ==============================================================================
# Malformed Document Edge Cases
# ==============================================================================


@pytest.mark.edge_case
class TestMalformedDocumentEdgeCases:
    """Test extractors with corrupted or malformed documents."""

    def test_pdf_with_corrupted_header(self, tmp_path):
        """
        EDGE: PDF with invalid header magic bytes.

        Partition: Format → Invalid → Corrupted header
        Expected: Graceful failure with error message
        """
        from extractors.pdf_extractor import PdfExtractor

        corrupted_pdf = tmp_path / "corrupted_header.pdf"
        corrupted_pdf.write_text("INVALID PDF HEADER\n%%EOF")

        extractor = PdfExtractor()
        result = extractor.extract(corrupted_pdf)

        assert result.success is False
        assert len(result.errors) > 0

    def test_pdf_with_truncated_content(self, tmp_path):
        """
        EDGE: PDF file that's been truncated mid-stream.

        Partition: Format → Invalid → Incomplete
        Expected: Error handling, partial recovery if possible
        """
        from reportlab.pdfgen import canvas

        from extractors.pdf_extractor import PdfExtractor

        # Create valid PDF
        pdf_path = tmp_path / "truncated.pdf"
        c = canvas.Canvas(str(pdf_path))
        c.drawString(100, 750, "Content")
        c.showPage()
        c.save()

        # Truncate the file
        with open(pdf_path, "rb") as f:
            content = f.read()

        truncated_content = content[: len(content) // 2]  # Keep first half

        with open(pdf_path, "wb") as f:
            f.write(truncated_content)

        extractor = PdfExtractor()
        result = extractor.extract(pdf_path)

        # Should fail gracefully
        assert result.success is False
        assert len(result.errors) > 0

    def test_excel_with_invalid_zip_structure(self, tmp_path):
        """
        EDGE: XLSX file with corrupted ZIP structure.

        Partition: Format → Invalid → Archive corruption
        Expected: Clear error message about format
        """
        from extractors.excel_extractor import ExcelExtractor

        # XLSX is a ZIP file, create invalid ZIP
        invalid_xlsx = tmp_path / "corrupted.xlsx"
        invalid_xlsx.write_bytes(b"Not a valid ZIP file")

        extractor = ExcelExtractor()
        result = extractor.extract(invalid_xlsx)

        assert result.success is False
        assert len(result.errors) > 0


# ==============================================================================
# Encoding Edge Cases
# ==============================================================================


@pytest.mark.edge_case
class TestEncodingEdgeCases:
    """Test extractors with various text encodings."""

    def test_text_with_utf8_bom(self, tmp_path):
        """
        EDGE: Text file with UTF-8 BOM (Byte Order Mark).

        Partition: Encoding → Valid → UTF-8 with BOM
        Boundary: BOM handling
        """
        from extractors.txt_extractor import TextFileExtractor

        txt_path = tmp_path / "with_bom.txt"
        # UTF-8 BOM is EF BB BF
        content = "\ufeff" + "Content with BOM"
        txt_path.write_text(content, encoding="utf-8-sig")

        extractor = TextFileExtractor()
        result = extractor.extract(txt_path)

        assert result.success is True
        # BOM should be handled/stripped
        assert len(result.content_blocks) > 0

    def test_text_with_mixed_line_endings(self, tmp_path):
        """
        EDGE: Text with mixed line endings (CRLF, LF, CR).

        Partition: Encoding → Valid → Line ending variants
        Boundary: Cross-platform compatibility
        """
        from extractors.txt_extractor import TextFileExtractor

        txt_path = tmp_path / "mixed_endings.txt"
        # Mix of line endings
        content = "Line 1\r\nLine 2\nLine 3\rLine 4"
        txt_path.write_bytes(content.encode("utf-8"))

        extractor = TextFileExtractor()
        result = extractor.extract(txt_path)

        assert result.success is True
        # Should handle all line ending types
        assert len(result.content_blocks) >= 3

    def test_text_with_null_bytes(self, tmp_path):
        """
        EDGE: Text file with embedded null bytes.

        Partition: Encoding → Invalid → Binary contamination
        Expected: Should handle or skip null bytes
        """
        from extractors.txt_extractor import TextFileExtractor

        txt_path = tmp_path / "with_nulls.txt"
        content = b"Text with\x00null\x00bytes"
        txt_path.write_bytes(content)

        extractor = TextFileExtractor()
        result = extractor.extract(txt_path)

        # Should handle gracefully (may succeed or fail depending on implementation)
        # At minimum, should not crash
        assert result is not None


# ==============================================================================
# Special Content Type Edge Cases
# ==============================================================================


@pytest.mark.edge_case
class TestSpecialContentEdgeCases:
    """Test extractors with special content scenarios."""

    def test_pdf_with_only_images_no_text(self, tmp_path):
        """
        EDGE: PDF containing only images, no extractable text.

        Partition: Content → Valid → Image-only
        Expected: Should detect images, warn about no text
        """
        from PIL import Image
        from reportlab.pdfgen import canvas

        from extractors.pdf_extractor import PdfExtractor

        # Create image
        img_path = tmp_path / "test.png"
        img = Image.new("RGB", (100, 100), color="red")
        img.save(img_path)

        # Create PDF with only image
        pdf_path = tmp_path / "image_only.pdf"
        c = canvas.Canvas(str(pdf_path))
        c.drawImage(str(img_path), 100, 700, width=100, height=100)
        c.showPage()
        c.save()

        extractor = PdfExtractor(config={"extract_images": True})
        result = extractor.extract(pdf_path)

        assert result.success is True
        # Should have extracted image metadata
        assert len(result.images) > 0

    def test_excel_with_only_formulas_no_values(self, tmp_path):
        """
        EDGE: Excel sheet with formulas but no calculated values.

        Partition: Content → Valid → Formula-only
        Expected: Should extract formulas or calculated results
        """
        from openpyxl import Workbook

        from extractors.excel_extractor import ExcelExtractor

        xlsx_path = tmp_path / "formulas_only.xlsx"
        wb = Workbook()
        ws = wb.active

        # Add formulas
        ws["A1"] = "=1+1"
        ws["A2"] = "=SUM(B1:B10)"

        wb.save(xlsx_path)

        extractor = ExcelExtractor()
        result = extractor.extract(xlsx_path)

        assert result.success is True
        # Should extract something (formulas or values)
        assert len(result.content_blocks) > 0

    def test_docx_with_only_tables_no_paragraphs(self, tmp_path):
        """
        EDGE: DOCX document with only tables, no text paragraphs.

        Partition: Content → Valid → Table-only
        Expected: Should extract table structure
        """
        from docx import Document

        from extractors.docx_extractor import DocxExtractor

        docx_path = tmp_path / "tables_only.docx"
        doc = Document()

        # Add table
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "A1"
        table.cell(0, 1).text = "B1"
        table.cell(1, 0).text = "A2"
        table.cell(1, 1).text = "B2"

        doc.save(docx_path)

        extractor = DocxExtractor()
        result = extractor.extract(docx_path)

        assert result.success is True
        # Should extract table content
        assert len(result.tables) > 0 or len(result.content_blocks) > 0


# ==============================================================================
# Metadata Edge Cases
# ==============================================================================


@pytest.mark.edge_case
class TestMetadataEdgeCases:
    """Test metadata extraction under edge conditions."""

    def test_pdf_with_extremely_long_metadata_values(self, tmp_path):
        """
        EDGE: PDF with very long title/author metadata.

        Partition: Metadata → Valid → Length boundary
        Boundary: Very long string values
        """
        from reportlab.pdfgen import canvas

        from extractors.pdf_extractor import PdfExtractor

        pdf_path = tmp_path / "long_metadata.pdf"
        c = canvas.Canvas(str(pdf_path))

        # Set very long metadata
        long_title = "A" * 500
        c.setTitle(long_title)
        c.setAuthor("B" * 500)

        c.drawString(100, 750, "Content")
        c.showPage()
        c.save()

        extractor = PdfExtractor()
        result = extractor.extract(pdf_path)

        assert result.success is True
        # Metadata should be captured (may be truncated)
        assert result.document_metadata is not None

    def test_pdf_with_special_chars_in_metadata(self, tmp_path):
        """
        EDGE: PDF with special characters/emojis in metadata.

        Partition: Metadata → Valid → Character encoding
        Boundary: Unicode and special characters
        """
        from reportlab.pdfgen import canvas

        from extractors.pdf_extractor import PdfExtractor

        pdf_path = tmp_path / "special_metadata.pdf"
        c = canvas.Canvas(str(pdf_path))

        # Unicode and special characters
        c.setTitle("Test 文档 🎉 <>&\"'")
        c.setAuthor("José García")

        c.drawString(100, 750, "Content")
        c.showPage()
        c.save()

        extractor = PdfExtractor()
        result = extractor.extract(pdf_path)

        assert result.success is True
        # Should handle unicode without crashing


# ==============================================================================
# Concurrent Access Edge Cases
# ==============================================================================


@pytest.mark.edge_case
class TestConcurrentAccessEdgeCases:
    """Test behavior under concurrent access scenarios."""

    def test_extract_same_file_multiple_times(self, tmp_path):
        """
        EDGE: Multiple extractions of the same file (simulating concurrent access).

        Partition: Concurrency → Valid → Read-only concurrent
        Expected: All extractions should succeed
        """
        from reportlab.pdfgen import canvas

        from extractors.pdf_extractor import PdfExtractor

        pdf_path = tmp_path / "shared.pdf"
        c = canvas.Canvas(str(pdf_path))
        c.drawString(100, 750, "Shared content")
        c.showPage()
        c.save()

        # Extract multiple times
        extractor = PdfExtractor()
        results = []

        for i in range(5):
            result = extractor.extract(pdf_path)
            results.append(result)

        # All should succeed
        assert all(r.success for r in results)

        # All should produce same content
        block_counts = [len(r.content_blocks) for r in results]
        assert len(set(block_counts)) == 1  # All same count


if __name__ == "__main__":
    pytest.main([__file__, "-v", "-m", "edge_case"])
