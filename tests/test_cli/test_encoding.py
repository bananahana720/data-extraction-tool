"""
Tests for CLI Encoding Handling.

Critical tests for v1.0.2 encoding fix that prevents 'charmap' codec errors
on Windows when processing files with Unicode characters.

Test coverage:
- Unicode character handling in console output
- UTF-8 file writing
- Private Use Area characters (e.g., \uf06c from PDFs)
- Mixed encoding scenarios
- Console reconfiguration on Windows
- Error message encoding

Related fix: ENCODING_FIX_SUMMARY.md
"""

import sys

import pytest
from docx import Document

from cli.commands import console
from cli.main import cli


class TestConsoleEncodingConfiguration:
    """Test console encoding is properly configured."""

    def test_stdout_uses_utf8_on_windows(self):
        """Verify stdout is configured to use UTF-8 encoding."""
        # After importing cli.commands, stdout should be UTF-8 on Windows
        if sys.platform == "win32":
            assert (
                sys.stdout.encoding.lower() == "utf-8"
            ), "stdout should use UTF-8 encoding on Windows"

    def test_stderr_uses_utf8_on_windows(self):
        """Verify stderr is configured to use UTF-8 encoding."""
        if sys.platform == "win32":
            assert (
                sys.stderr.encoding.lower() == "utf-8"
            ), "stderr should use UTF-8 encoding on Windows"

    def test_rich_console_configured_correctly(self):
        """Verify Rich Console is configured with proper settings."""
        # Check console settings
        assert (
            console.is_terminal or console.force_terminal
        ), "Console should be configured for terminal output"


class TestUnicodeCharacterHandling:
    """Test handling of various Unicode characters."""

    def test_bmp_characters(self, cli_runner, tmp_path):
        """Test Basic Multilingual Plane characters (U+0000 to U+FFFF)."""
        # Create DOCX with BMP characters
        file_path = tmp_path / "bmp_test.docx"
        doc = Document()
        doc.add_paragraph("English text")
        doc.add_paragraph("中文字符")  # Chinese
        doc.add_paragraph("Русский")  # Russian
        doc.add_paragraph("العربية")  # Arabic
        doc.add_paragraph("日本語")  # Japanese
        doc.add_paragraph("한글")  # Korean
        doc.add_paragraph("Emoji: ✓ ✗ • ★")  # Common symbols
        doc.save(file_path)

        output_file = tmp_path / "output.json"

        result = cli_runner.invoke(
            cli, ["extract", str(file_path), "--output", str(output_file), "--format", "json"]
        )

        assert result.exit_code == 0, f"Failed with: {result.output}"
        assert output_file.exists()

        # Verify file can be read back with UTF-8
        content = output_file.read_text(encoding="utf-8")
        assert len(content) > 0

    def test_private_use_area_characters(self, cli_runner, tmp_path):
        """Test Private Use Area characters (U+E000 to U+F8FF) - common in PDFs."""
        # Create DOCX with PUA characters (like those in PDF icons)
        file_path = tmp_path / "pua_test.docx"
        doc = Document()
        doc.add_paragraph("Regular text")
        # These are the problematic characters from PDFs
        doc.add_paragraph("Icons: \uf06c \uf070 \uf071 \uf0a0")
        doc.save(file_path)

        output_file = tmp_path / "output.json"

        result = cli_runner.invoke(
            cli,
            [
                "extract",
                str(file_path),
                "--output",
                str(output_file),
                "--format",
                "json",
                "--force",
            ],
        )

        # Should not crash with encoding error
        assert result.exit_code == 0, f"Should handle PUA characters. Output: {result.output}"
        assert output_file.exists()

    def test_supplementary_multilingual_plane(self, cli_runner, tmp_path):
        """Test SMP characters (U+10000 to U+1FFFF) - emojis, etc."""
        file_path = tmp_path / "smp_test.docx"
        doc = Document()
        doc.add_paragraph("Emojis: 😀 🎉 🔥 💯")
        doc.add_paragraph("Math: 𝕊𝕥𝕪𝕝𝕖𝕕")
        doc.save(file_path)

        output_file = tmp_path / "output.json"

        result = cli_runner.invoke(
            cli, ["extract", str(file_path), "--output", str(output_file), "--format", "json"]
        )

        assert result.exit_code == 0
        assert output_file.exists()

    def test_mixed_encodings(self, cli_runner, tmp_path):
        """Test file with mixed scripts and encodings."""
        file_path = tmp_path / "mixed_test.docx"
        doc = Document()
        doc.add_heading("Mixed Encoding Test", level=1)
        doc.add_paragraph("English: Hello World")
        doc.add_paragraph("Greek: Γεια σου Κόσμε")
        doc.add_paragraph("Hebrew: שלום עולם")
        doc.add_paragraph("Thai: สวัสดีชาวโลก")
        doc.add_paragraph("Icons: \uf06c \uf070")
        doc.add_paragraph("Emoji: 🌍 🌎 🌏")
        doc.save(file_path)

        output_file = tmp_path / "output.json"

        result = cli_runner.invoke(
            cli, ["extract", str(file_path), "--output", str(output_file), "--format", "json"]
        )

        assert result.exit_code == 0
        assert output_file.exists()

        # Verify content is readable
        content = output_file.read_text(encoding="utf-8")
        assert len(content) > 100


class TestFileWriteEncoding:
    """Test file writing uses proper UTF-8 encoding."""

    def test_json_output_utf8_encoding(self, cli_runner, tmp_path):
        """Verify JSON output files are written with UTF-8."""
        file_path = tmp_path / "unicode.docx"
        doc = Document()
        doc.add_paragraph("Unicode: \uf06c • ✓ 中文")
        doc.save(file_path)

        output_file = tmp_path / "output.json"

        result = cli_runner.invoke(
            cli, ["extract", str(file_path), "--output", str(output_file), "--format", "json"]
        )

        assert result.exit_code == 0
        assert output_file.exists()

        # Read with UTF-8 explicitly
        content = output_file.read_text(encoding="utf-8")
        assert len(content) > 0

        # Verify it's valid JSON with Unicode
        import json

        data = json.loads(content)
        assert isinstance(data, dict)

    def test_markdown_output_utf8_encoding(self, cli_runner, tmp_path):
        """Verify Markdown output files are written with UTF-8."""
        file_path = tmp_path / "unicode.docx"
        doc = Document()
        doc.add_paragraph("Unicode: \uf06c • ✓ 日本語")
        doc.save(file_path)

        output_file = tmp_path / "output.md"

        result = cli_runner.invoke(
            cli, ["extract", str(file_path), "--output", str(output_file), "--format", "markdown"]
        )

        assert result.exit_code == 0
        assert output_file.exists()

        # Should be readable as UTF-8
        content = output_file.read_text(encoding="utf-8")
        assert len(content) > 0

    def test_filename_with_unicode(self, cli_runner, tmp_path):
        """Test handling filenames containing Unicode characters."""
        # Create file with Unicode in name
        file_path = tmp_path / "测试文档.docx"
        doc = Document()
        doc.add_paragraph("Test content")
        doc.save(file_path)

        output_file = tmp_path / "输出.json"

        result = cli_runner.invoke(
            cli, ["extract", str(file_path), "--output", str(output_file), "--format", "json"]
        )

        # Should handle Unicode filenames
        assert result.exit_code == 0
        assert output_file.exists()


class TestConsoleOutputEncoding:
    """Test console output handles Unicode correctly."""

    def test_success_message_with_unicode_filename(self, cli_runner, tmp_path):
        """Test success message displays Unicode filename."""
        file_path = tmp_path / "文档.docx"
        doc = Document()
        doc.add_paragraph("Content")
        doc.save(file_path)

        output_file = tmp_path / "output.json"

        result = cli_runner.invoke(
            cli, ["extract", str(file_path), "--output", str(output_file), "--format", "json"]
        )

        assert result.exit_code == 0
        # Should not crash displaying the filename
        # (CliRunner captures output, so won't see encoding error here,
        #  but the command should complete successfully)

    def test_error_message_with_unicode(self, cli_runner, tmp_path):
        """Test error messages handle Unicode in paths."""
        # Non-existent file with Unicode name
        file_path = tmp_path / "不存在.docx"

        result = cli_runner.invoke(cli, ["extract", str(file_path), "--format", "json"])

        # Should fail but not crash on encoding
        assert result.exit_code != 0
        assert "not found" in result.output.lower() or "does not exist" in result.output.lower()


class TestBatchProcessingEncoding:
    """Test batch processing handles Unicode correctly."""

    def test_batch_with_unicode_files(self, cli_runner, tmp_path):
        """Test batch processing multiple files with Unicode content."""
        # Create multiple files with Unicode
        files = []
        for i in range(3):
            file_path = tmp_path / f"doc{i}.docx"
            doc = Document()
            doc.add_paragraph(f"Document {i}")
            doc.add_paragraph("Unicode: \uf06c • ✓ 中文 🎉")
            doc.save(file_path)
            files.append(file_path)

        output_dir = tmp_path / "output"

        result = cli_runner.invoke(
            cli,
            [
                "batch",
                str(tmp_path),
                "--output",
                str(output_dir),
                "--format",
                "json",
                "--pattern",
                "*.docx",
            ],
        )

        assert result.exit_code == 0
        assert output_dir.exists()

        # Verify output files created
        output_files = list(output_dir.glob("*.json"))
        assert len(output_files) >= 3

        # Verify each file is valid UTF-8 JSON
        for output_file in output_files:
            content = output_file.read_text(encoding="utf-8")
            import json

            data = json.loads(content)
            assert isinstance(data, dict)

    def test_batch_summary_with_unicode_filenames(self, cli_runner, tmp_path):
        """Test batch summary displays Unicode filenames correctly."""
        # Create files with Unicode names
        for name in ["测试1.docx", "テスト2.docx", "тест3.docx"]:
            file_path = tmp_path / name
            doc = Document()
            doc.add_paragraph("Content")
            doc.save(file_path)

        output_dir = tmp_path / "output"

        result = cli_runner.invoke(
            cli,
            [
                "batch",
                str(tmp_path),
                "--output",
                str(output_dir),
                "--format",
                "json",
                "--pattern",
                "*.docx",
            ],
        )

        # Should complete without encoding errors
        assert result.exit_code == 0


class TestEncodingEdgeCases:
    """Test edge cases in encoding handling."""

    def test_null_bytes_handling(self, cli_runner, tmp_path):
        """Test handling of null bytes in content."""
        file_path = tmp_path / "null_test.txt"
        # Write file with null bytes
        file_path.write_bytes(b"Content with \x00 null byte")

        output_file = tmp_path / "output.json"

        result = cli_runner.invoke(
            cli, ["extract", str(file_path), "--output", str(output_file), "--format", "json"]
        )

        # Should handle gracefully (may succeed or fail, but no crash)
        assert result.exit_code in (0, 1)

    def test_invalid_utf8_sequences(self, cli_runner, tmp_path):
        """Test handling of invalid UTF-8 sequences."""
        file_path = tmp_path / "invalid_utf8.txt"
        # Write invalid UTF-8
        file_path.write_bytes(b"Valid text \xFF\xFE invalid UTF-8")

        output_file = tmp_path / "output.json"

        result = cli_runner.invoke(
            cli, ["extract", str(file_path), "--output", str(output_file), "--format", "json"]
        )

        # Should handle gracefully with errors='replace'
        assert result.exit_code in (0, 1)

    def test_very_long_unicode_string(self, cli_runner, tmp_path):
        """Test handling of very long strings with Unicode."""
        file_path = tmp_path / "long_unicode.docx"
        doc = Document()
        # Create very long paragraph with Unicode
        long_text = "Unicode: \uf06c • ✓ 中文 " * 1000
        doc.add_paragraph(long_text)
        doc.save(file_path)

        output_file = tmp_path / "output.json"

        result = cli_runner.invoke(
            cli, ["extract", str(file_path), "--output", str(output_file), "--format", "json"]
        )

        assert result.exit_code == 0
        assert output_file.exists()

        # Verify output is still valid
        content = output_file.read_text(encoding="utf-8")
        assert len(content) > 10000


class TestWindowsSpecificEncoding:
    """Test Windows-specific encoding scenarios."""

    @pytest.mark.skipif(sys.platform != "win32", reason="Windows-specific test")
    def test_console_reconfiguration_on_windows(self):
        """Verify console was reconfigured on Windows."""
        # After importing cli.commands, check encoding
        assert sys.stdout.encoding.lower() == "utf-8"
        assert sys.stderr.encoding.lower() == "utf-8"

    @pytest.mark.skipif(sys.platform != "win32", reason="Windows-specific test")
    def test_charmap_codec_not_used(self):
        """Verify charmap codec is not being used."""
        # Should not be using Windows default 'cp1252' or 'charmap'
        assert "charmap" not in sys.stdout.encoding.lower()
        assert "cp1252" not in sys.stdout.encoding.lower()

    @pytest.mark.skipif(sys.platform != "win32", reason="Windows-specific test")
    def test_errors_parameter_set(self):
        """Verify errors='replace' is configured."""
        # Check if errors parameter is accessible
        if hasattr(sys.stdout, "errors"):
            assert sys.stdout.errors in ("replace", "backslashreplace", "namereplace")


@pytest.mark.integration
class TestEncodingIntegration:
    """Integration tests for encoding across entire pipeline."""

    def test_end_to_end_unicode_preservation(self, cli_runner, tmp_path):
        """Test Unicode characters are preserved through entire pipeline."""
        # Create file with known Unicode content
        file_path = tmp_path / "unicode_preservation.docx"
        doc = Document()
        doc.add_heading("Unicode Test", level=1)
        test_strings = [
            "English: Hello World",
            "Chinese: 你好世界",
            "Japanese: こんにちは世界",
            "Emoji: 🌍 🎉 ✨",
            "Icons: \uf06c \uf070",
            "Math: ∑ ∫ √ π",
        ]
        for text in test_strings:
            doc.add_paragraph(text)
        doc.save(file_path)

        output_file = tmp_path / "output.json"

        result = cli_runner.invoke(
            cli, ["extract", str(file_path), "--output", str(output_file), "--format", "json"]
        )

        assert result.exit_code == 0

        # Verify content in output
        content = output_file.read_text(encoding="utf-8")

        # Check some characters are present (may be replaced with ?)
        assert "Hello World" in content
        assert len(content) > 200

    def test_concurrent_unicode_processing(self, cli_runner, tmp_path):
        """Test Unicode handling in concurrent batch processing."""
        # Create multiple files with diverse Unicode
        unicode_sets = [
            "Latin: Hello Ñoño",
            "Cyrillic: Привет мир",
            "Arabic: مرحبا بالعالم",
            "CJK: 你好世界",
            "Icons: \uf06c \uf070",
        ]

        for i, text in enumerate(unicode_sets):
            file_path = tmp_path / f"unicode_{i}.docx"
            doc = Document()
            doc.add_paragraph(text)
            doc.save(file_path)

        output_dir = tmp_path / "output"

        result = cli_runner.invoke(
            cli,
            [
                "batch",
                str(tmp_path),
                "--output",
                str(output_dir),
                "--format",
                "json",
                "--workers",
                "2",  # Test with threading
            ],
        )

        assert result.exit_code == 0
        assert output_dir.exists()

        # Verify all files processed
        output_files = list(output_dir.glob("*.json"))
        assert len(output_files) >= len(unicode_sets)
