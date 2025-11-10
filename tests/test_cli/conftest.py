"""
Test fixtures for CLI tests.

Provides reusable fixtures for CLI testing including:
- Sample document files (DOCX, PDF, PPTX, XLSX)
- Configured pipeline instances
- Click test runner
- Temporary output directories
"""

import pytest
from pathlib import Path
from click.testing import CliRunner
from docx import Document

from src.pipeline import ExtractionPipeline, BatchProcessor
from src.extractors import DocxExtractor
from src.processors import ContextLinker, MetadataAggregator
from src.formatters import JsonFormatter, MarkdownFormatter
from src.infrastructure import ConfigManager


@pytest.fixture
def cli_runner():
    """
    Provide Click test runner for CLI testing.

    Returns:
        CliRunner instance for invoking CLI commands
    """
    return CliRunner()


@pytest.fixture
def sample_docx_file(tmp_path):
    """
    Create a sample DOCX file for testing.

    Args:
        tmp_path: pytest temporary directory fixture

    Returns:
        Path to created DOCX file
    """
    file_path = tmp_path / "sample.docx"

    # Create simple DOCX with content
    doc = Document()
    doc.add_heading("Test Document", level=1)
    doc.add_paragraph("This is a test paragraph.")
    doc.add_paragraph("This is another paragraph.")
    doc.save(file_path)

    return file_path


@pytest.fixture
def sample_text_file(tmp_path):
    """
    Create a sample text file for testing.

    Args:
        tmp_path: pytest temporary directory fixture

    Returns:
        Path to created text file
    """
    file_path = tmp_path / "sample.txt"
    file_path.write_text("Line 1\nLine 2\nLine 3\n")
    return file_path


@pytest.fixture
def multiple_test_files(tmp_path):
    """
    Create multiple test files for batch processing.

    Args:
        tmp_path: pytest temporary directory fixture

    Returns:
        List of paths to created files
    """
    files = []

    # Create 3 DOCX files
    for i in range(1, 4):
        file_path = tmp_path / f"doc{i}.docx"
        doc = Document()
        doc.add_heading(f"Document {i}", level=1)
        doc.add_paragraph(f"Content for document {i}")
        doc.save(file_path)
        files.append(file_path)

    # Create 2 text files
    for i in range(1, 3):
        file_path = tmp_path / f"text{i}.txt"
        file_path.write_text(f"Text file {i} content\n")
        files.append(file_path)

    return files


@pytest.fixture
def configured_pipeline():
    """
    Create a fully configured extraction pipeline for testing.

    Returns:
        ExtractionPipeline with extractors, processors, and formatters
    """
    pipeline = ExtractionPipeline()

    # Register extractors
    pipeline.register_extractor("docx", DocxExtractor())

    # Add processors
    pipeline.add_processor(ContextLinker())
    pipeline.add_processor(MetadataAggregator())

    # Add formatters
    pipeline.add_formatter(JsonFormatter())
    pipeline.add_formatter(MarkdownFormatter())

    return pipeline


@pytest.fixture
def configured_batch_processor(configured_pipeline):
    """
    Create a configured batch processor for testing.

    Args:
        configured_pipeline: Pipeline fixture

    Returns:
        BatchProcessor instance
    """
    return BatchProcessor(pipeline=configured_pipeline, max_workers=2)  # Small number for testing


@pytest.fixture
def config_file(tmp_path):
    """
    Create a sample configuration file for testing.

    Args:
        tmp_path: pytest temporary directory fixture

    Returns:
        Path to config file
    """
    import yaml

    config_path = tmp_path / "config.yaml"
    config_data = {
        "pipeline": {
            "max_workers": 4,
            "timeout_per_file": 300,
        },
        "extractors": {
            "docx": {
                "skip_empty": True,
            },
        },
        "output": {
            "default_format": "json",
        },
    }

    with open(config_path, "w") as f:
        yaml.dump(config_data, f)

    return config_path


@pytest.fixture
def output_directory(tmp_path):
    """
    Create a temporary output directory for testing.

    Args:
        tmp_path: pytest temporary directory fixture

    Returns:
        Path to output directory
    """
    output_dir = tmp_path / "output"
    output_dir.mkdir()
    return output_dir


@pytest.fixture
def nonexistent_file(tmp_path):
    """
    Provide path to a file that doesn't exist for error testing.

    Args:
        tmp_path: pytest temporary directory fixture

    Returns:
        Path to nonexistent file
    """
    return tmp_path / "nonexistent.docx"


@pytest.fixture
def unsupported_file(tmp_path):
    """
    Create a file with unsupported format for error testing.

    Args:
        tmp_path: pytest temporary directory fixture

    Returns:
        Path to unsupported file
    """
    file_path = tmp_path / "unsupported.xyz"
    file_path.write_text("This is an unsupported format")
    return file_path
