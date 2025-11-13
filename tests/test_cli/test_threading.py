"""
Tests for CLI Threading and Concurrency.

Critical tests for v1.0.2 threading fix that prevents deadlocks in batch
processing with progress display updates from worker threads.

Test coverage:
- Thread-safe progress updates
- Concurrent file processing
- Worker thread management
- Thread cleanup
- Exception handling in threads
- Lock contention
- Progress display under load

Related fix: BATCH_STALLING_FIX.md
"""

import threading
import time
from concurrent.futures import ThreadPoolExecutor

import pytest
from click.testing import CliRunner
from docx import Document

from cli.main import cli
from cli.progress_display import BatchProgress, SingleFileProgress


class TestProgressDisplayThreadSafety:
    """Test progress display classes are thread-safe."""

    def test_single_file_progress_has_lock(self, tmp_path):
        """Verify SingleFileProgress has threading lock."""
        file_path = tmp_path / "test.docx"
        progress = SingleFileProgress(file_path=file_path, quiet=True)

        assert hasattr(progress, "_lock")
        assert isinstance(progress._lock, threading.Lock)

    def test_batch_progress_has_lock(self, tmp_path):
        """Verify BatchProgress has threading lock."""
        files = [tmp_path / f"test{i}.docx" for i in range(3)]
        progress = BatchProgress(file_paths=files, quiet=True)

        assert hasattr(progress, "_lock")
        assert isinstance(progress._lock, threading.Lock)

    def test_single_file_progress_concurrent_updates(self, tmp_path):
        """Test SingleFileProgress handles concurrent updates."""
        file_path = tmp_path / "test.docx"

        with SingleFileProgress(file_path=file_path, quiet=False) as progress:
            # Simulate multiple threads updating progress
            def update_progress(thread_id):
                for i in range(10):
                    progress.update(
                        {
                            "stage": f"thread_{thread_id}",
                            "percentage": i * 10,
                            "message": f"Update from thread {thread_id}",
                        }
                    )
                    time.sleep(0.001)  # Small delay

            threads = []
            for i in range(5):
                t = threading.Thread(target=update_progress, args=(i,))
                threads.append(t)
                t.start()

            # Wait for all threads
            for t in threads:
                t.join(timeout=5)

            # Should complete without deadlock
            assert all(not t.is_alive() for t in threads)

    def test_batch_progress_concurrent_updates(self, tmp_path):
        """Test BatchProgress handles concurrent updates."""
        files = [tmp_path / f"test{i}.docx" for i in range(5)]

        with BatchProgress(file_paths=files, quiet=False) as progress:
            # Simulate worker threads updating progress
            def update_file_progress(file_path, thread_id):
                for i in range(10):
                    progress.update(
                        {
                            "current_file": file_path.name,
                            "percentage": i * 10,
                            "stage": f"processing_{i}",
                            "items_processed": thread_id,
                        }
                    )
                    time.sleep(0.001)

                progress.mark_file_complete(file_path, success=True)

            threads = []
            for i, file_path in enumerate(files):
                t = threading.Thread(target=update_file_progress, args=(file_path, i))
                threads.append(t)
                t.start()

            # Wait for all threads
            for t in threads:
                t.join(timeout=5)

            # Should complete without deadlock
            assert all(not t.is_alive() for t in threads)

    def test_progress_update_exception_handling(self, tmp_path):
        """Test progress updates handle exceptions gracefully."""
        file_path = tmp_path / "test.docx"

        with SingleFileProgress(file_path=file_path, quiet=False) as progress:
            # Try to cause exception by passing invalid status
            def bad_update():
                for i in range(5):
                    try:
                        progress.update(None)  # Invalid status
                    except:
                        pass  # Should be caught internally
                    time.sleep(0.001)

            threads = [threading.Thread(target=bad_update) for _ in range(3)]
            for t in threads:
                t.start()
            for t in threads:
                t.join(timeout=2)

            # Should not deadlock even with exceptions
            assert all(not t.is_alive() for t in threads)


class TestBatchWorkerManagement:
    """Test batch processing worker thread management."""

    def test_batch_with_single_worker(self, cli_runner, multiple_test_files, tmp_path):
        """Test batch processing with single worker (no threading)."""
        input_dir = multiple_test_files[0].parent
        output_dir = tmp_path / "output"

        result = cli_runner.invoke(
            cli,
            [
                "batch",
                str(input_dir),
                "--output",
                str(output_dir),
                "--format",
                "json",
                "--workers",
                "1",
            ],
        )

        assert result.exit_code == 0
        assert output_dir.exists()

    def test_batch_with_multiple_workers(self, cli_runner, multiple_test_files, tmp_path):
        """Test batch processing with multiple workers."""
        input_dir = multiple_test_files[0].parent
        output_dir = tmp_path / "output"

        for worker_count in [2, 4, 8]:
            output_subdir = output_dir / f"workers_{worker_count}"

            result = cli_runner.invoke(
                cli,
                [
                    "batch",
                    str(input_dir),
                    "--output",
                    str(output_subdir),
                    "--format",
                    "json",
                    "--workers",
                    str(worker_count),
                ],
            )

            assert result.exit_code == 0, f"Failed with {worker_count} workers: {result.output}"
            assert output_subdir.exists()

    def test_batch_worker_count_validation(self, cli_runner, tmp_path):
        """Test worker count validation."""
        input_dir = tmp_path
        output_dir = tmp_path / "output"

        # Zero workers should fail
        result = cli_runner.invoke(
            cli, ["batch", str(input_dir), "--output", str(output_dir), "--workers", "0"]
        )
        assert result.exit_code != 0

        # Negative workers should fail
        result = cli_runner.invoke(
            cli, ["batch", str(input_dir), "--output", str(output_dir), "--workers", "-1"]
        )
        assert result.exit_code != 0

    def test_batch_many_files_with_workers(self, cli_runner, tmp_path):
        """Test batch processing many files with multiple workers."""
        # Create 20 test files
        for i in range(20):
            file_path = tmp_path / f"doc{i}.docx"
            doc = Document()
            doc.add_paragraph(f"Document {i} content")
            doc.save(file_path)

        output_dir = tmp_path / "output"

        result = cli_runner.invoke(
            cli,
            [
                "batch",
                str(tmp_path),
                "--output",
                str(output_dir),
                "--format",
                "json",
                "--pattern",
                "*.docx",
                "--workers",
                "4",
            ],
        )

        assert result.exit_code == 0
        assert output_dir.exists()

        # Verify all files processed
        output_files = list(output_dir.glob("*.json"))
        assert len(output_files) >= 20


class TestThreadCleanup:
    """Test proper cleanup of threads."""

    def test_threads_cleaned_up_after_batch(self, cli_runner, multiple_test_files, tmp_path):
        """Verify threads are cleaned up after batch processing."""
        input_dir = multiple_test_files[0].parent
        output_dir = tmp_path / "output"

        # Count threads before
        threads_before = threading.active_count()

        result = cli_runner.invoke(
            cli,
            [
                "batch",
                str(input_dir),
                "--output",
                str(output_dir),
                "--format",
                "json",
                "--workers",
                "4",
            ],
        )

        assert result.exit_code == 0

        # Give threads time to cleanup
        time.sleep(0.5)

        # Count threads after
        threads_after = threading.active_count()

        # Should not have leaked threads (allow some tolerance)
        assert (
            threads_after <= threads_before + 2
        ), f"Thread leak detected: {threads_before} before, {threads_after} after"

    def test_no_zombie_threads(self, cli_runner, multiple_test_files, tmp_path):
        """Verify no zombie threads remain."""
        input_dir = multiple_test_files[0].parent
        output_dir = tmp_path / "output"

        result = cli_runner.invoke(
            cli,
            [
                "batch",
                str(input_dir),
                "--output",
                str(output_dir),
                "--format",
                "json",
                "--workers",
                "4",
            ],
        )

        assert result.exit_code == 0

        # Wait for cleanup
        time.sleep(0.5)

        # Check for daemon threads (should be minimal)
        active_threads = threading.enumerate()
        daemon_threads = [t for t in active_threads if t.daemon]

        # Most threads should be cleaned up
        assert len(daemon_threads) < 10


class TestExceptionHandlingInThreads:
    """Test exception handling in worker threads."""

    def test_single_file_failure_doesnt_crash_batch(self, cli_runner, tmp_path):
        """Test that one file failure doesn't crash entire batch."""
        # Create valid files
        for i in range(3):
            file_path = tmp_path / f"valid{i}.docx"
            doc = Document()
            doc.add_paragraph(f"Valid document {i}")
            doc.save(file_path)

        # Create corrupted file
        corrupted = tmp_path / "corrupted.docx"
        corrupted.write_bytes(b"This is not a valid DOCX file")

        output_dir = tmp_path / "output"

        result = cli_runner.invoke(
            cli,
            [
                "batch",
                str(tmp_path),
                "--output",
                str(output_dir),
                "--format",
                "json",
                "--pattern",
                "*.docx",
                "--workers",
                "2",
            ],
        )

        # Should complete but may have errors
        assert result.exit_code in (0, 1)  # May fail due to corrupted file
        assert output_dir.exists()

        # Valid files should still be processed
        output_files = list(output_dir.glob("*.json"))
        assert len(output_files) >= 3

    def test_thread_exception_isolation(self, tmp_path):
        """Test that exception in one thread doesn't affect others."""
        files = [tmp_path / f"test{i}.docx" for i in range(5)]

        with BatchProgress(file_paths=files, quiet=True) as progress:
            results = {"success": 0, "failed": 0}
            lock = threading.Lock()

            def process_file(file_path, should_fail):
                try:
                    if should_fail:
                        raise ValueError(f"Simulated error for {file_path.name}")

                    # Simulate processing
                    for i in range(10):
                        progress.update({"current_file": file_path.name, "percentage": i * 10})
                        time.sleep(0.001)

                    progress.mark_file_complete(file_path, success=True)
                    with lock:
                        results["success"] += 1
                except Exception as e:
                    progress.mark_file_failed(file_path, str(e))
                    with lock:
                        results["failed"] += 1

            # Create threads, some will fail
            threads = []
            for i, file_path in enumerate(files):
                should_fail = i % 2 == 0  # Fail every other file
                t = threading.Thread(target=process_file, args=(file_path, should_fail))
                threads.append(t)
                t.start()

            # Wait for all threads
            for t in threads:
                t.join(timeout=5)

            # All threads should complete
            assert all(not t.is_alive() for t in threads)
            assert results["success"] + results["failed"] == len(files)


class TestLockContention:
    """Test behavior under high lock contention."""

    def test_high_contention_updates(self, tmp_path):
        """Test progress updates under high contention."""
        file_path = tmp_path / "test.docx"

        with SingleFileProgress(file_path=file_path, quiet=False) as progress:
            update_count = {"count": 0}
            lock = threading.Lock()

            def rapid_updates(thread_id):
                for i in range(100):
                    progress.update(
                        {
                            "stage": f"t{thread_id}",
                            "percentage": (i % 100),
                        }
                    )
                    with lock:
                        update_count["count"] += 1

            # Many threads doing rapid updates
            threads = [threading.Thread(target=rapid_updates, args=(i,)) for i in range(10)]

            for t in threads:
                t.start()

            for t in threads:
                t.join(timeout=10)

            # All threads should complete
            assert all(not t.is_alive() for t in threads)
            # Should have attempted all updates
            assert update_count["count"] == 1000

    def test_lock_timeout_behavior(self, tmp_path):
        """Test that lock doesn't cause indefinite blocking."""
        files = [tmp_path / f"test{i}.docx" for i in range(10)]

        with BatchProgress(file_paths=files, quiet=False) as progress:
            completed = {"count": 0}
            lock = threading.Lock()

            def blocking_update(file_path):
                start_time = time.time()

                # Try many updates
                for i in range(50):
                    progress.update({"current_file": file_path.name, "percentage": i * 2})

                # Should complete within reasonable time
                elapsed = time.time() - start_time
                assert elapsed < 5, f"Updates took too long: {elapsed}s"

                with lock:
                    completed["count"] += 1

            threads = [threading.Thread(target=blocking_update, args=(f,)) for f in files]

            for t in threads:
                t.start()

            for t in threads:
                t.join(timeout=10)

            # All should complete
            assert completed["count"] == len(files)


class TestConcurrentFileProcessing:
    """Test concurrent processing of files."""

    def test_concurrent_extract_doesnt_conflict(self, cli_runner, tmp_path):
        """Test that concurrent extractions don't conflict."""
        # Create test files
        files = []
        for i in range(5):
            file_path = tmp_path / f"doc{i}.docx"
            doc = Document()
            doc.add_paragraph(f"Document {i}")
            doc.save(file_path)
            files.append(file_path)

        # Process concurrently using ThreadPoolExecutor
        def extract_file(file_path):
            output = file_path.with_suffix(".json")
            runner = CliRunner()
            result = runner.invoke(
                cli,
                ["extract", str(file_path), "--output", str(output), "--format", "json", "--quiet"],
            )
            return result.exit_code == 0

        with ThreadPoolExecutor(max_workers=3) as executor:
            results = list(executor.map(extract_file, files))

        # All should succeed
        assert all(results)
        # All output files should exist
        assert all(f.with_suffix(".json").exists() for f in files)

    def test_no_race_conditions_in_output(self, cli_runner, tmp_path):
        """Test no race conditions when writing outputs."""
        # Create multiple files
        for i in range(10):
            file_path = tmp_path / f"doc{i}.docx"
            doc = Document()
            doc.add_paragraph(f"Unique content for document {i}")
            doc.save(file_path)

        output_dir = tmp_path / "output"

        result = cli_runner.invoke(
            cli,
            [
                "batch",
                str(tmp_path),
                "--output",
                str(output_dir),
                "--format",
                "json",
                "--pattern",
                "*.docx",
                "--workers",
                "4",
            ],
        )

        assert result.exit_code == 0

        # Verify each output file has unique content
        output_files = list(output_dir.glob("*.json"))
        assert len(output_files) >= 10

        # Each file should be valid and unique
        contents = []
        for output_file in output_files:
            content = output_file.read_text(encoding="utf-8")
            assert content not in contents, "Duplicate content detected!"
            contents.append(content)


@pytest.mark.stress
class TestThreadingStress:
    """Stress tests for threading behavior."""

    def test_many_workers_stress(self, cli_runner, tmp_path):
        """Test with many workers (stress test)."""
        # Create many small files
        for i in range(50):
            file_path = tmp_path / f"doc{i}.docx"
            doc = Document()
            doc.add_paragraph(f"Doc {i}")
            doc.save(file_path)

        output_dir = tmp_path / "output"

        # Test with many workers
        result = cli_runner.invoke(
            cli,
            [
                "batch",
                str(tmp_path),
                "--output",
                str(output_dir),
                "--format",
                "json",
                "--pattern",
                "*.docx",
                "--workers",
                "16",
            ],
        )

        assert result.exit_code == 0
        output_files = list(output_dir.glob("*.json"))
        assert len(output_files) >= 50

    def test_rapid_batch_processing(self, cli_runner, tmp_path):
        """Test rapid successive batch operations."""
        # Create test files
        for i in range(10):
            file_path = tmp_path / f"doc{i}.docx"
            doc = Document()
            doc.add_paragraph(f"Document {i}")
            doc.save(file_path)

        # Run batch multiple times rapidly
        for run in range(3):
            output_dir = tmp_path / f"output_{run}"

            result = cli_runner.invoke(
                cli,
                [
                    "batch",
                    str(tmp_path),
                    "--output",
                    str(output_dir),
                    "--format",
                    "json",
                    "--pattern",
                    "*.docx",
                    "--workers",
                    "4",
                ],
            )

            assert result.exit_code == 0
            assert output_dir.exists()

            # Small delay between runs
            time.sleep(0.1)


@pytest.mark.integration
class TestThreadingIntegration:
    """Integration tests for threading."""

    def test_full_batch_workflow_with_threading(self, cli_runner, tmp_path):
        """Test complete batch workflow with threading."""
        # Create diverse files
        for i in range(15):
            file_path = tmp_path / f"doc{i}.docx"
            doc = Document()
            doc.add_heading(f"Document {i}", level=1)
            for j in range(5):
                doc.add_paragraph(f"Paragraph {j} of document {i}")
            doc.save(file_path)

        output_dir = tmp_path / "output"

        # Process with threading
        result = cli_runner.invoke(
            cli,
            [
                "batch",
                str(tmp_path),
                "--output",
                str(output_dir),
                "--format",
                "all",  # All formats
                "--pattern",
                "*.docx",
                "--workers",
                "4",
                "--verbose",  # Verbose mode with threading
            ],
        )

        assert result.exit_code == 0
        assert output_dir.exists()

        # Verify outputs
        json_files = list(output_dir.glob("*.json"))
        md_files = list(output_dir.glob("*.md"))

        assert len(json_files) >= 15
        assert len(md_files) >= 15
