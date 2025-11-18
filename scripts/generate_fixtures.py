#!/usr/bin/env python3
"""
Fixture Data Generator

Generates test fixtures including sample documents (PDF, DOCX, XLSX),
semantic corpus documents, PII-free test data, and edge case scenarios.

Usage:
    python scripts/generate_fixtures.py --types pdf docx xlsx
    python scripts/generate_fixtures.py --semantic-corpus --count 10
    python scripts/generate_fixtures.py --edge-cases --output-dir tests/fixtures/edge_cases
"""

import argparse
import hashlib
import json
import random
import sys
from datetime import datetime, timedelta
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

import structlog  # type: ignore[import-not-found]
import yaml  # type: ignore[import-untyped]

# Document generation libraries
try:
    from reportlab.lib.pagesizes import letter  # type: ignore[import-untyped]
    from reportlab.lib.units import inch  # type: ignore[import-untyped]
    from reportlab.pdfgen import canvas  # type: ignore[import-untyped]

    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

try:
    from docx import Document  # type: ignore[import-not-found]
    from docx.shared import Pt  # type: ignore[import-not-found]

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from openpyxl import Workbook  # type: ignore[import-untyped]
    from openpyxl.styles import Font, PatternFill  # type: ignore[import-untyped]

    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False

# Configure structured logging
logger = structlog.get_logger()

# Constants
PROJECT_ROOT = Path(__file__).parent.parent
FIXTURES_DIR = PROJECT_ROOT / "tests" / "fixtures"
GENERATED_DIR = FIXTURES_DIR / "generated"

# Deterministic seed for reproducibility
DEFAULT_SEED = 42

# Sample data pools (PII-free)
FIRST_NAMES = [
    "Alice",
    "Bob",
    "Charlie",
    "Diana",
    "Edward",
    "Fiona",
    "George",
    "Helen",
    "Ivan",
    "Julia",
]
LAST_NAMES = [
    "Smith",
    "Johnson",
    "Williams",
    "Brown",
    "Jones",
    "Garcia",
    "Miller",
    "Davis",
    "Rodriguez",
    "Martinez",
]
COMPANIES = [
    "Acme Corp",
    "TechCo",
    "GlobalTech",
    "InnovateSoft",
    "DataDynamics",
    "CloudFirst",
    "AgileWorks",
    "NextGen Systems",
    "FutureTech",
    "SmartSolutions",
]
DEPARTMENTS = [
    "Engineering",
    "Sales",
    "Marketing",
    "Finance",
    "HR",
    "Operations",
    "IT",
    "Legal",
    "R&D",
    "Customer Service",
]
PRODUCTS = [
    "Widget Pro",
    "DataAnalyzer",
    "CloudManager",
    "SecureVault",
    "AutoProcessor",
    "SmartMonitor",
    "QuickSync",
    "PowerTool",
    "MegaApp",
    "UltraService",
]
LOREM_WORDS = [
    "lorem",
    "ipsum",
    "dolor",
    "sit",
    "amet",
    "consectetur",
    "adipiscing",
    "elit",
    "sed",
    "do",
    "eiusmod",
    "tempor",
    "incididunt",
    "ut",
    "labore",
    "et",
    "dolore",
    "magna",
    "aliqua",
]


class FixtureGenerator:
    """Generates various types of test fixtures with deterministic output."""

    def __init__(
        self,
        output_dir: Path = GENERATED_DIR,
        seed: int = DEFAULT_SEED,
        config_file: Optional[Path] = None,
    ):
        """
        Initialize the fixture generator.

        Args:
            output_dir: Directory for output fixtures
            seed: Random seed for reproducibility
            config_file: Optional configuration file
        """
        self.output_dir = output_dir
        self.seed = seed
        self.config = self._load_config(config_file) if config_file else {}

        # Set seed for deterministic output
        random.seed(seed)

        # Create output directories
        self.output_dir.mkdir(parents=True, exist_ok=True)
        (self.output_dir / "documents").mkdir(exist_ok=True)
        (self.output_dir / "semantic").mkdir(exist_ok=True)
        (self.output_dir / "edge_cases").mkdir(exist_ok=True)

        logger.info("initialized_fixture_generator", output_dir=str(output_dir), seed=seed)

    def _load_config(self, config_file: Path) -> Dict[str, Any]:
        """Load configuration from YAML or JSON file."""
        if config_file.suffix in [".yaml", ".yml"]:
            with open(config_file) as f:
                result = yaml.safe_load(f)
                return result if isinstance(result, dict) else {}
        elif config_file.suffix == ".json":
            with open(config_file) as f:
                result = json.load(f)
                return result if isinstance(result, dict) else {}
        else:
            logger.warning("unsupported_config_format", file=str(config_file))
            return {}

    def generate_pdf(
        self,
        filename: str = "sample.pdf",
        pages: int = 1,
        content_type: str = "standard",
    ) -> Path:
        """
        Generate a PDF document with configurable content.

        Args:
            filename: Output filename
            pages: Number of pages
            content_type: Type of content (standard, financial, technical, legal)

        Returns:
            Path to generated PDF
        """
        if not REPORTLAB_AVAILABLE:
            logger.error("reportlab_not_available")
            raise ImportError("reportlab is required for PDF generation")

        output_path = self.output_dir / "documents" / filename

        c = canvas.Canvas(str(output_path), pagesize=letter)
        width, height = letter

        for page_num in range(1, pages + 1):
            # Add page header
            c.setFont("Helvetica-Bold", 16)
            c.drawString(
                1 * inch, height - 1 * inch, f"{content_type.title()} Document - Page {page_num}"
            )

            # Generate content based on type
            c.setFont("Helvetica", 12)
            y_position = height - 1.5 * inch

            if content_type == "financial":
                content = self._generate_financial_content()
            elif content_type == "technical":
                content = self._generate_technical_content()
            elif content_type == "legal":
                content = self._generate_legal_content()
            else:
                content = self._generate_standard_content()

            # Write content lines
            for line in content[:30]:  # Max 30 lines per page
                if y_position < 1 * inch:
                    break
                c.drawString(1 * inch, y_position, line)
                y_position -= 20

            # Add page footer
            c.setFont("Helvetica", 10)
            c.drawString(width / 2 - 50, 0.5 * inch, f"Page {page_num} of {pages}")

            if page_num < pages:
                c.showPage()

        c.save()
        logger.info("pdf_generated", path=str(output_path), pages=pages, type=content_type)
        return output_path

    def generate_docx(
        self,
        filename: str = "sample.docx",
        paragraphs: int = 5,
        content_type: str = "standard",
        with_tables: bool = False,
        with_images: bool = False,
    ) -> Path:
        """
        Generate a DOCX document with configurable content.

        Args:
            filename: Output filename
            paragraphs: Number of paragraphs
            content_type: Type of content
            with_tables: Include tables
            with_images: Include placeholder images

        Returns:
            Path to generated DOCX
        """
        if not DOCX_AVAILABLE:
            logger.error("python_docx_not_available")
            raise ImportError("python-docx is required for DOCX generation")

        output_path = self.output_dir / "documents" / filename

        doc = Document()

        # Add title
        doc.add_heading(f"{content_type.title()} Document", 0)

        # Add metadata as subtitle
        doc.add_heading(f"Generated: {datetime.now().strftime('%Y-%m-%d')}", level=2)

        # Generate content paragraphs
        for i in range(paragraphs):
            if content_type == "financial":
                content = self._generate_financial_content()
            elif content_type == "technical":
                content = self._generate_technical_content()
            else:
                content = self._generate_standard_content()

            p = doc.add_paragraph(" ".join(content[i * 3 : (i + 1) * 3]))

            # Add some formatting variation
            if i % 2 == 0:
                run = p.runs[0]
                run.font.size = Pt(11)

        # Add table if requested
        if with_tables:
            doc.add_heading("Data Table", level=2)
            table = doc.add_table(rows=5, cols=4)
            table.style = "Light Grid Accent 1"

            # Add headers
            headers = ["ID", "Name", "Department", "Value"]
            for i, header in enumerate(headers):
                table.rows[0].cells[i].text = header

            # Add data
            for row_idx in range(1, 5):
                table.rows[row_idx].cells[0].text = str(1000 + row_idx)
                table.rows[row_idx].cells[
                    1
                ].text = f"{random.choice(FIRST_NAMES)} {random.choice(LAST_NAMES)}"
                table.rows[row_idx].cells[2].text = random.choice(DEPARTMENTS)
                table.rows[row_idx].cells[3].text = f"${random.randint(1000, 99999):,.2f}"

        # Save document
        doc.save(str(output_path))
        logger.info(
            "docx_generated", path=str(output_path), paragraphs=paragraphs, type=content_type
        )
        return output_path

    def generate_xlsx(
        self,
        filename: str = "sample.xlsx",
        sheets: int = 1,
        rows: int = 100,
        content_type: str = "standard",
    ) -> Path:
        """
        Generate an XLSX document with configurable content.

        Args:
            filename: Output filename
            sheets: Number of sheets
            rows: Number of rows per sheet
            content_type: Type of content (standard, financial, inventory)

        Returns:
            Path to generated XLSX
        """
        if not OPENPYXL_AVAILABLE:
            logger.error("openpyxl_not_available")
            raise ImportError("openpyxl is required for XLSX generation")

        output_path = self.output_dir / "documents" / filename

        wb = Workbook()

        for sheet_num in range(sheets):
            if sheet_num == 0:
                ws = wb.active
                ws.title = f"Sheet_{sheet_num + 1}"
            else:
                ws = wb.create_sheet(f"Sheet_{sheet_num + 1}")

            # Style headers
            header_font = Font(bold=True, size=12)
            header_fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")

            if content_type == "financial":
                headers = ["Date", "Account", "Description", "Debit", "Credit", "Balance"]
            elif content_type == "inventory":
                headers = ["SKU", "Product", "Category", "Quantity", "Unit Price", "Total Value"]
            else:
                headers = ["ID", "Name", "Department", "Email", "Phone", "Start Date"]

            # Write headers
            for col, header in enumerate(headers, 1):
                cell = ws.cell(row=1, column=col, value=header)
                cell.font = header_font
                cell.fill = header_fill

            # Generate data rows
            for row_idx in range(2, rows + 2):
                if content_type == "financial":
                    row_data = self._generate_financial_row()
                elif content_type == "inventory":
                    row_data = self._generate_inventory_row()
                else:
                    row_data = self._generate_standard_row()

                for col, value in enumerate(row_data, 1):
                    ws.cell(row=row_idx, column=col, value=value)

            # Adjust column widths
            for column in ws.columns:
                max_length = 0
                column_letter = column[0].column_letter
                for cell in column:
                    try:
                        if len(str(cell.value)) > max_length:
                            max_length = len(str(cell.value))
                    except (TypeError, AttributeError):
                        pass
                adjusted_width = min(max_length + 2, 50)
                ws.column_dimensions[column_letter].width = adjusted_width

        # Save workbook
        wb.save(str(output_path))
        logger.info(
            "xlsx_generated", path=str(output_path), sheets=sheets, rows=rows, type=content_type
        )
        return output_path

    def generate_semantic_corpus(
        self,
        count: int = 10,
        topics: Optional[List[str]] = None,
        min_words: int = 100,
        max_words: int = 500,
    ) -> List[Path]:
        """
        Generate documents with specific semantic patterns for NLP testing.

        Args:
            count: Number of documents to generate
            topics: List of topics/themes
            min_words: Minimum words per document
            max_words: Maximum words per document

        Returns:
            List of paths to generated documents
        """
        if topics is None:
            topics = ["technology", "finance", "healthcare", "education", "environment"]

        corpus_dir = self.output_dir / "semantic" / "corpus"
        corpus_dir.mkdir(parents=True, exist_ok=True)

        generated_files = []

        for i in range(count):
            topic = topics[i % len(topics)]
            word_count = random.randint(min_words, max_words)

            # Generate document with semantic coherence
            content = self._generate_semantic_document(topic, word_count)

            # Create filename with hash for determinism
            content_hash = hashlib.md5(content.encode()).hexdigest()[:8]
            filename = f"{topic}_{i:03d}_{content_hash}.txt"
            file_path = corpus_dir / filename

            file_path.write_text(content)
            generated_files.append(file_path)

            logger.info(
                "semantic_document_generated", path=str(file_path), topic=topic, words=word_count
            )

        return generated_files

    def _generate_semantic_document(self, topic: str, word_count: int) -> str:
        """Generate a semantically coherent document for a given topic."""
        topic_vocabularies = {
            "technology": [
                "software",
                "algorithm",
                "data",
                "system",
                "application",
                "interface",
                "cloud",
                "security",
                "network",
                "protocol",
                "API",
                "database",
                "framework",
                "deployment",
                "architecture",
            ],
            "finance": [
                "investment",
                "portfolio",
                "asset",
                "equity",
                "bond",
                "market",
                "trading",
                "analysis",
                "risk",
                "return",
                "dividend",
                "capital",
                "liquidity",
                "valuation",
                "strategy",
            ],
            "healthcare": [
                "patient",
                "treatment",
                "diagnosis",
                "medical",
                "clinical",
                "therapy",
                "medication",
                "symptom",
                "disease",
                "prevention",
                "care",
                "health",
                "hospital",
                "research",
                "protocol",
            ],
            "education": [
                "student",
                "curriculum",
                "learning",
                "teaching",
                "assessment",
                "course",
                "program",
                "skill",
                "knowledge",
                "training",
                "development",
                "academic",
                "instruction",
                "evaluation",
                "methodology",
            ],
            "environment": [
                "climate",
                "sustainability",
                "ecosystem",
                "conservation",
                "renewable",
                "pollution",
                "biodiversity",
                "habitat",
                "carbon",
                "emission",
                "resource",
                "energy",
                "waste",
                "recycling",
                "protection",
            ],
        }

        vocab = topic_vocabularies.get(topic, LOREM_WORDS)
        common_words = [
            "the",
            "and",
            "of",
            "to",
            "in",
            "for",
            "with",
            "on",
            "at",
            "by",
            "from",
            "about",
            "through",
            "during",
            "after",
        ]

        sentences = []
        words_written = 0

        while words_written < word_count:
            # Generate a sentence
            sentence_length = random.randint(8, 20)
            sentence_words = []

            for _ in range(sentence_length):
                if random.random() < 0.3:  # 30% topic-specific words
                    sentence_words.append(random.choice(vocab))
                elif random.random() < 0.7:  # 40% common words
                    sentence_words.append(random.choice(common_words))
                else:  # 30% lorem words for variety
                    sentence_words.append(random.choice(LOREM_WORDS))

            sentence = " ".join(sentence_words).capitalize() + "."
            sentences.append(sentence)
            words_written += sentence_length

        # Organize into paragraphs
        paragraphs = []
        sentences_per_paragraph = 3

        for i in range(0, len(sentences), sentences_per_paragraph):
            paragraph = " ".join(sentences[i : i + sentences_per_paragraph])
            paragraphs.append(paragraph)

        return "\n\n".join(paragraphs)

    def generate_edge_cases(self) -> Dict[str, Path]:
        """
        Generate documents with edge cases for testing.

        Returns:
            Dictionary mapping edge case type to file path
        """
        edge_cases = {}
        edge_dir = self.output_dir / "edge_cases"
        edge_dir.mkdir(parents=True, exist_ok=True)

        # Empty file
        empty_path = edge_dir / "empty.txt"
        empty_path.write_text("")
        edge_cases["empty"] = empty_path

        # Single character file
        single_char = edge_dir / "single_char.txt"
        single_char.write_text("X")
        edge_cases["single_char"] = single_char

        # Very large file (1MB of text)
        large_content = "X" * (1024 * 1024)  # 1MB
        large_path = edge_dir / "large_1mb.txt"
        large_path.write_text(large_content)
        edge_cases["large"] = large_path

        # Special characters
        special_chars = "!@#$%^&*(){}[]|\\:;\"'<>,.?/~`±§€£¥₹©®™"
        special_path = edge_dir / "special_chars.txt"
        special_path.write_text(special_chars)
        edge_cases["special_chars"] = special_path

        # Unicode characters
        unicode_content = "Hello 世界 مرحبا мир שלום κόσμος"
        unicode_path = edge_dir / "unicode.txt"
        unicode_path.write_text(unicode_content)
        edge_cases["unicode"] = unicode_path

        # Very long single line
        long_line = "word " * 10000  # 50,000 chars
        long_line_path = edge_dir / "long_line.txt"
        long_line_path.write_text(long_line)
        edge_cases["long_line"] = long_line_path

        # Many short lines
        many_lines = "\n".join(["Line " + str(i) for i in range(10000)])
        many_lines_path = edge_dir / "many_lines.txt"
        many_lines_path.write_text(many_lines)
        edge_cases["many_lines"] = many_lines_path

        # Mixed encodings (simulate)
        mixed_path = edge_dir / "mixed_encoding.txt"
        mixed_path.write_bytes(b"ASCII text\n\xc3\xa9 UTF-8\n\xff\xfe Unicode")
        edge_cases["mixed_encoding"] = mixed_path

        # Whitespace only
        whitespace = "   \t\n\r\n   \t   \n\n\n"
        whitespace_path = edge_dir / "whitespace_only.txt"
        whitespace_path.write_text(whitespace)
        edge_cases["whitespace"] = whitespace_path

        # Null bytes (binary-like)
        null_content = "Text\x00with\x00null\x00bytes"
        null_path = edge_dir / "null_bytes.txt"
        null_path.write_text(null_content)
        edge_cases["null_bytes"] = null_path

        logger.info("edge_cases_generated", count=len(edge_cases))
        return edge_cases

    def _generate_financial_content(self) -> List[str]:
        """Generate financial-themed content lines."""
        templates = [
            f"Q{random.randint(1,4)} revenue increased by {random.randint(5, 25)}% year-over-year.",
            f"Operating margin improved to {random.randint(15, 35)}% from previous quarter.",
            f"Total assets valued at ${random.randint(100, 999)},{random.randint(100, 999)},{random.randint(100, 999)}.",
            f"Net income for the period: ${random.randint(10, 99)}.{random.randint(0, 9)} million.",
            f"Earnings per share: ${random.randint(1, 9)}.{random.randint(10, 99)}.",
            f"Dividend declared: ${random.randint(0, 2)}.{random.randint(10, 99)} per share.",
        ]
        return [random.choice(templates) for _ in range(20)]

    def _generate_technical_content(self) -> List[str]:
        """Generate technical documentation content."""
        templates = [
            f"The {random.choice(['API', 'system', 'service'])} supports {random.choice(['REST', 'GraphQL', 'gRPC'])} protocol.",
            f"Maximum throughput: {random.randint(1000, 10000)} requests per second.",
            f"Latency: p99 < {random.randint(10, 100)}ms, p50 < {random.randint(5, 50)}ms.",
            f"Database: {random.choice(['PostgreSQL', 'MySQL', 'MongoDB', 'Redis'])} v{random.randint(10, 15)}.{random.randint(0, 9)}.",
            f"Memory usage: {random.randint(100, 999)}MB average, {random.randint(1, 4)}GB peak.",
            f"CPU utilization: {random.randint(20, 80)}% under normal load.",
        ]
        return [random.choice(templates) for _ in range(20)]

    def _generate_legal_content(self) -> List[str]:
        """Generate legal document content."""
        templates = [
            "WHEREAS, the parties have agreed to the following terms and conditions:",
            "The agreement shall commence on the Effective Date and continue for a period of twelve (12) months.",
            "Neither party shall disclose Confidential Information without prior written consent.",
            "This Agreement shall be governed by the laws of the applicable jurisdiction.",
            "IN WITNESS WHEREOF, the parties have executed this Agreement as of the date last written below.",
            "Force Majeure: Neither party shall be liable for delays caused by circumstances beyond reasonable control.",
        ]
        return templates * 3 + [f"Section {i+1}. {random.choice(templates)}" for i in range(5)]

    def _generate_standard_content(self) -> List[str]:
        """Generate standard lorem ipsum style content."""
        lines = []
        for _ in range(20):
            words = random.sample(LOREM_WORDS, random.randint(5, 15))
            lines.append(" ".join(words).capitalize() + ".")
        return lines

    def _generate_financial_row(self) -> Tuple:
        """Generate a financial data row."""
        date = datetime.now() - timedelta(days=random.randint(1, 365))
        return (
            date.strftime("%Y-%m-%d"),
            f"ACC-{random.randint(1000, 9999)}",
            random.choice(["Payment", "Invoice", "Transfer", "Adjustment"]),
            random.randint(0, 10000) if random.random() > 0.5 else 0,
            random.randint(0, 10000) if random.random() > 0.5 else 0,
            random.randint(10000, 99999),
        )

    def _generate_inventory_row(self) -> Tuple:
        """Generate an inventory data row."""
        return (
            f"SKU-{random.randint(10000, 99999)}",
            random.choice(PRODUCTS),
            random.choice(["Electronics", "Software", "Hardware", "Accessories"]),
            random.randint(0, 1000),
            round(random.uniform(9.99, 999.99), 2),
            round(random.uniform(999.99, 99999.99), 2),
        )

    def _generate_standard_row(self) -> Tuple:
        """Generate a standard data row."""
        first_name = random.choice(FIRST_NAMES)
        last_name = random.choice(LAST_NAMES)
        email = f"{first_name.lower()}.{last_name.lower()}@example.com"
        date = datetime.now() - timedelta(days=random.randint(1, 1825))

        return (
            random.randint(1000, 9999),
            f"{first_name} {last_name}",
            random.choice(DEPARTMENTS),
            email,
            f"+1-555-{random.randint(100, 999)}-{random.randint(1000, 9999)}",
            date.strftime("%Y-%m-%d"),
        )

    def validate_fixtures(self) -> Dict[str, bool]:
        """
        Validate that generated fixtures are PII-free and properly formatted.

        Returns:
            Dictionary of validation results
        """
        results = {}

        # Check for common PII patterns
        pii_patterns = [
            r"\b\d{3}-\d{2}-\d{4}\b",  # SSN
            r"\b\d{16}\b",  # Credit card
            r"\b[A-Z]{2}\d{6,8}\b",  # Passport
        ]

        import re

        for file_path in self.output_dir.rglob("*"):
            if file_path.is_file():
                try:
                    if file_path.suffix in [".txt", ".csv"]:
                        content = file_path.read_text()
                        has_pii = any(re.search(pattern, content) for pattern in pii_patterns)
                        results[str(file_path)] = not has_pii
                    else:
                        results[str(file_path)] = True  # Assume binary files are OK
                except Exception as e:
                    logger.error("validation_error", file=str(file_path), error=str(e))
                    results[str(file_path)] = False

        return results


def parse_arguments() -> argparse.Namespace:
    """Parse command line arguments."""
    parser = argparse.ArgumentParser(
        description="Generate test fixtures for data extraction testing",
        formatter_class=argparse.ArgumentDefaultsHelpFormatter,
    )

    parser.add_argument(
        "--output-dir",
        type=Path,
        default=GENERATED_DIR,
        help="Output directory for fixtures",
    )

    parser.add_argument(
        "--config",
        type=Path,
        help="Configuration file (YAML or JSON)",
    )

    parser.add_argument(
        "--seed",
        type=int,
        default=DEFAULT_SEED,
        help="Random seed for deterministic output",
    )

    # Document generation options
    parser.add_argument(
        "--types",
        nargs="+",
        choices=["pdf", "docx", "xlsx", "all"],
        help="Document types to generate",
    )

    parser.add_argument(
        "--count",
        type=int,
        default=1,
        help="Number of documents per type",
    )

    parser.add_argument(
        "--content-type",
        choices=["standard", "financial", "technical", "legal", "inventory"],
        default="standard",
        help="Type of content to generate",
    )

    # Semantic corpus options
    parser.add_argument(
        "--semantic-corpus",
        action="store_true",
        help="Generate semantic corpus documents",
    )

    parser.add_argument(
        "--corpus-count",
        type=int,
        default=10,
        help="Number of corpus documents",
    )

    parser.add_argument(
        "--topics",
        nargs="+",
        help="Topics for semantic corpus",
    )

    # Edge cases
    parser.add_argument(
        "--edge-cases",
        action="store_true",
        help="Generate edge case documents",
    )

    # Validation
    parser.add_argument(
        "--validate",
        action="store_true",
        help="Validate fixtures for PII",
    )

    parser.add_argument(
        "--dry-run",
        action="store_true",
        help="Preview without generating files",
    )

    parser.add_argument(
        "--verbose",
        action="store_true",
        help="Enable verbose logging",
    )

    return parser.parse_args()


def main() -> int:
    """Main entry point."""
    try:
        args = parse_arguments()

        # Configure logging
        if args.verbose:
            import logging

            structlog.configure(
                wrapper_class=structlog.make_filtering_bound_logger(logging.DEBUG),
            )

        if args.dry_run:
            print("=== DRY RUN MODE ===")
            print(f"Would generate fixtures in: {args.output_dir}")
            print(f"Seed: {args.seed}")
            if args.types:
                print(f"Document types: {args.types}")
                print(f"Count per type: {args.count}")
                print(f"Content type: {args.content_type}")
            if args.semantic_corpus:
                print(f"Semantic corpus: {args.corpus_count} documents")
                print(f"Topics: {args.topics or 'default'}")
            if args.edge_cases:
                print("Edge cases: 10 types")
            return 0

        # Initialize generator
        generator = FixtureGenerator(
            output_dir=args.output_dir,
            seed=args.seed,
            config_file=args.config,
        )

        # Generate document fixtures
        if args.types:
            types = args.types if "all" not in args.types else ["pdf", "docx", "xlsx"]

            for doc_type in types:
                for i in range(args.count):
                    filename = f"{args.content_type}_{i+1:03d}.{doc_type}"

                    if doc_type == "pdf":
                        generator.generate_pdf(
                            filename=filename,
                            pages=random.randint(1, 5),
                            content_type=args.content_type,
                        )
                    elif doc_type == "docx":
                        generator.generate_docx(
                            filename=filename,
                            paragraphs=random.randint(3, 10),
                            content_type=args.content_type,
                            with_tables=args.content_type in ["financial", "inventory"],
                        )
                    elif doc_type == "xlsx":
                        generator.generate_xlsx(
                            filename=filename,
                            sheets=random.randint(1, 3),
                            rows=random.randint(50, 200),
                            content_type=args.content_type,
                        )

            print(f"✅ Generated {len(types) * args.count} documents")

        # Generate semantic corpus
        if args.semantic_corpus:
            files = generator.generate_semantic_corpus(
                count=args.corpus_count,
                topics=args.topics,
            )
            print(f"✅ Generated {len(files)} semantic corpus documents")

        # Generate edge cases
        if args.edge_cases:
            edge_cases = generator.generate_edge_cases()
            print(f"✅ Generated {len(edge_cases)} edge case documents")
            for case_type, path in edge_cases.items():
                print(f"   - {case_type}: {path.name}")

        # Validate fixtures
        if args.validate:
            results = generator.validate_fixtures()
            valid_count = sum(1 for v in results.values() if v)
            print(f"\n📊 Validation Results: {valid_count}/{len(results)} files are PII-free")

            invalid_files = [f for f, valid in results.items() if not valid]
            if invalid_files:
                print("⚠️  Files with potential PII:")
                for f in invalid_files:
                    print(f"   - {f}")

        return 0

    except Exception as e:
        logger.error("fixture_generation_failed", error=str(e))
        print(f"❌ Error: {e}", file=sys.stderr)
        return 1


if __name__ == "__main__":
    sys.exit(main())
